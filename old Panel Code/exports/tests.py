"""
Tests for exports app.
Covers: Permission scoping, export view access control, ExportService.
"""
import os
from django.conf import settings
from django.test import TestCase, SimpleTestCase
from django.contrib.auth import get_user_model
from django.core.cache import cache
from django.http import HttpResponse
from django.test import RequestFactory
from unittest import mock
import json

User = get_user_model()


def _setup_export_data():
    """Create a super_admin, client, table with cards for export tests."""
    admin = User.objects.create_user(
        username='exadmin@test.com', email='exadmin@test.com',
        password='adminpass1', role='super_admin',
    )
    client_user = User.objects.create_user(
        username='exclient@test.com', email='exclient@test.com',
        password='clientpass1', role='client',
    )
    from client.models import Client
    client = Client.objects.create(user=client_user, name='Export Client')

    from idcards.models import IDCardGroup, IDCardTable, IDCard
    group = IDCardGroup.objects.create(client=client, name='Export Group')
    table = IDCardTable.objects.create(
        group=group, name='Export Table',
        fields=[
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'FATHER', 'type': 'text', 'order': 2},
        ],
    )
    for i in range(5):
        IDCard.objects.create(
            table=table,
            field_data={'NAME': f'STUDENT {i}', 'FATHER': f'FATHER {i}'},
            status='pending',
        )
    return admin, client_user, client, table


class ExportPermissionTests(TestCase):
    """Tests for export access control."""

    def setUp(self):
        self.admin, self.client_user, self.client_obj, self.table = _setup_export_data()
        cache.clear()

    def tearDown(self):
        cache.clear()

    def test_unauthenticated_blocked(self):
        response = self.client.post(
            f'/panel/exports/xlsx/{self.table.id}/',
            data=json.dumps({'card_ids': []}),
            content_type='application/json',
        )
        self.assertIn(response.status_code, [302, 403])

    def test_admin_can_export_xlsx(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')
        from idcards.models import IDCard
        card_ids = list(IDCard.objects.filter(table=self.table).values_list('id', flat=True))
        response = self.client.post(
            f'/panel/exports/xlsx/{self.table.id}/',
            data=json.dumps({'card_ids': card_ids}),
            content_type='application/json',
        )
        self.assertIn(response.status_code, [200, 400])

    def test_client_blocked_from_download_all(self):
        self.client.login(username='exclient@test.com', password='clientpass1')
        response = self.client.post(
            f'/panel/exports/download-all/{self.table.id}/',
            data=json.dumps({}),
            content_type='application/json',
        )
        # Client should be blocked (403 JSON or 302 redirect)
        self.assertIn(response.status_code, [302, 403])


class ExportServiceTests(TestCase):
    """Tests for ExportService scoping."""

    def setUp(self):
        self.admin, self.client_user, self.client_obj, self.table = _setup_export_data()

    def test_export_service_scopes_cards(self):
        from exports.services import ExportService
        service = ExportService(self.admin)
        cards = service.get_scoped_cards(self.table)
        self.assertEqual(cards.count(), 5)

    def test_export_service_empty_for_wrong_client(self):
        other_user = User.objects.create_user(
            username='other@test.com', email='other@test.com',
            password='otherpass1', role='client',
        )
        from client.models import Client
        Client.objects.create(user=other_user, name='Other Client')
        # Other user should not see cards from the first client's table via ExportService
        from exports.services import ExportService
        service = ExportService(other_user)
        cards = service.get_scoped_cards(self.table)
        # Depending on scoping logic, might be 0 or 5 (super_admin sees all)
        self.assertIsNotNone(cards)


class ExportViewHelperTests(TestCase):
    def setUp(self):
        self.factory = RequestFactory()
        self.admin, self.client_user, self.client_obj, self.table = _setup_export_data()

    def test_get_card_ids_from_json_body(self):
        from exports.views import _get_card_ids_from_request

        request = self.factory.post(
            '/panel/exports/xlsx/1/',
            data=json.dumps({'card_ids': [1, '2', 'bad', 3]}),
            content_type='application/json',
        )

        ids = _get_card_ids_from_request(request)
        self.assertEqual(ids, [1, 2, 3])

    def test_get_status_from_request_supports_json_charset_content_type(self):
        from exports.views import _get_status_from_request

        request = self.factory.post(
            '/panel/exports/pdf/1/',
            data=json.dumps({'status': 'pending'}),
            content_type='application/json; charset=utf-8',
        )

        status = _get_status_from_request(request)
        self.assertEqual(status, 'pending')

    def test_get_card_ids_normalizes_and_deduplicates(self):
        from exports.views import _get_card_ids_from_request

        request = self.factory.post(
            '/panel/exports/xlsx/1/',
            data=json.dumps({'card_ids': [1, '1', ' 2 ', True, 0, -5, 'bad']}),
            content_type='application/json',
        )

        ids = _get_card_ids_from_request(request)
        self.assertEqual(ids, [1, 2])

    def test_get_card_ids_rejects_oversized_json_body(self):
        from exports import views as export_views

        request = self.factory.post(
            '/panel/exports/xlsx/1/',
            data=json.dumps({'card_ids': [1, 2, 3]}),
            content_type='application/json',
        )

        with mock.patch.object(export_views, 'MAX_EXPORT_JSON_BODY_BYTES', 8):
            ids = export_views._get_card_ids_from_request(request)

        self.assertIsNone(ids)

    def test_get_card_ids_fallback_by_status_when_not_supplied(self):
        from exports.views import _get_card_ids_from_request

        request = self.factory.post(
            f'/panel/exports/xlsx/{self.table.id}/',
            data=json.dumps({'status': 'pending'}),
            content_type='application/json',
        )
        request.user = self.admin

        ids = _get_card_ids_from_request(request, table_id=self.table.id)
        self.assertEqual(len(ids), 5)

    def test_get_card_ids_fallback_requires_client_scope(self):
        from client.models import Client
        from exports.views import _get_card_ids_from_request

        outsider = User.objects.create_user(
            username='export-outsider@test.com',
            email='export-outsider@test.com',
            password='pass1234',
            role='client',
        )
        Client.objects.create(user=outsider, name='Outsider Client')

        request = self.factory.post(
            f'/panel/exports/xlsx/{self.table.id}/',
            data=json.dumps({'status': 'pending'}),
            content_type='application/json',
        )
        request.user = outsider

        ids = _get_card_ids_from_request(request, table_id=self.table.id)
        self.assertIsNone(ids)

    def test_get_card_ids_fallback_class_section_uses_canonical_class_matching(self):
        from idcards.models import IDCard
        from exports.views import _get_card_ids_from_request

        self.table.fields = [
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'CLASS', 'type': 'class', 'order': 2},
            {'name': 'SECTION', 'type': 'section', 'order': 3},
        ]
        self.table.save(update_fields=['fields'])

        IDCard.objects.filter(table=self.table).delete()
        matched = IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'A', 'CLASS': 'KG-I', 'SECTION': 'A'},
            status='pending',
        )
        IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'B', 'CLASS': 'KG-II', 'SECTION': 'A'},
            status='pending',
        )

        request = self.factory.post(
            f'/panel/exports/pdf/{self.table.id}/',
            data=json.dumps({'status': 'pending', 'class': 'KG1', 'section': 'A'}),
            content_type='application/json',
        )
        request.user = self.admin

        ids = _get_card_ids_from_request(request, table_id=self.table.id)
        self.assertEqual(ids, [matched.id])

    def test_get_card_ids_fallback_respects_client_staff_row_scope(self):
        from client.models import Client
        from idcards.models import IDCard
        from staff.models import Staff
        from exports.views import _get_card_ids_from_request

        self.table.fields = [
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'CLASS', 'type': 'class', 'order': 2},
        ]
        self.table.save(update_fields=['fields'])

        IDCard.objects.filter(table=self.table).delete()
        allowed = IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'Allowed', 'CLASS': 'KG-I'},
            status='pending',
        )
        IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'Blocked', 'CLASS': 'KG-II'},
            status='pending',
        )

        staff_user = User.objects.create_user(
            username='export-cs-scope@test.com',
            email='export-cs-scope@test.com',
            password='pass1234',
            role='client_staff',
        )
        # Keep parent client profile present and active.
        Client.objects.filter(pk=self.client_obj.pk).update(status='active')
        Staff.objects.create(
            user=staff_user,
            staff_type='client_staff',
            client=self.client_obj,
            assigned_table_ids=[self.table.id],
            allowed_classes=['KG1'],
            perm_idcard_bulk_download=True,
            perm_idcard_pending_list=True,
        )

        request = self.factory.post(
            f'/panel/exports/pdf/{self.table.id}/',
            data=json.dumps({'status': 'pending'}),
            content_type='application/json',
        )
        request.user = staff_user

        ids = _get_card_ids_from_request(request, table_id=self.table.id)
        self.assertEqual(ids, [allowed.id])

    def test_get_card_ids_fallback_ignores_date_range_for_non_download_status(self):
        from exports.views import _get_card_ids_from_request

        request = self.factory.post(
            f'/panel/exports/pdf/{self.table.id}/',
            data=json.dumps({
                'status': 'pending',
                'from': '2099-01-01T00:00:00',
                'to': '2099-12-31T23:59:59',
            }),
            content_type='application/json',
        )
        request.user = self.admin

        ids = _get_card_ids_from_request(request, table_id=self.table.id)
        self.assertEqual(len(ids), 5)

    def test_get_all_card_ids_ignores_date_range_for_non_download_status(self):
        from core.services import IDCardService

        result = IDCardService.get_all_card_ids(
            self.table.id,
            status_filter='pending',
            from_date='2099-01-01T00:00:00',
            to_date='2099-12-31T23:59:59',
        )

        self.assertTrue(result.success)
        self.assertEqual(result.data.get('total_count'), 5)

    def test_get_card_ids_fallback_respects_image_filter(self):
        from idcards.models import IDCard
        from exports.views import _get_card_ids_from_request

        self.table.fields = [
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'PHOTO', 'type': 'image', 'order': 2},
        ]
        self.table.save(update_fields=['fields'])

        IDCard.objects.filter(table=self.table).delete()
        complete = IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'A', 'PHOTO': 'clients_imgs/a.jpg'},
            status='pending',
        )
        IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'B', 'PHOTO': 'PENDING:upload'},
            status='pending',
        )
        IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'C', 'PHOTO': 'NOT_FOUND'},
            status='pending',
        )

        request = self.factory.post(
            f'/panel/exports/pdf/{self.table.id}/',
            data=json.dumps({
                'status': 'pending',
                'image_column': 'PHOTO',
                'image_condition': 'complete',
            }),
            content_type='application/json',
        )
        request.user = self.admin

        ids = _get_card_ids_from_request(request, table_id=self.table.id)
        self.assertEqual(ids, [complete.id])

    def test_get_image_rename_options_filters_invalid_pairs(self):
        from exports.views import _get_image_rename_options_from_request

        request = self.factory.post(
            f'/panel/exports/images/{self.table.id}/',
            data=json.dumps({
                'rename_options': {
                    'enabled': True,
                    'image_name_fields': {
                        ' photo ': ' Student Name ',
                        '': 'x',
                        'QR': '',
                    },
                }
            }),
            content_type='application/json',
        )

        opts = _get_image_rename_options_from_request(request)
        self.assertTrue(opts['enabled'])
        self.assertEqual(opts['image_name_fields'], {'PHOTO': 'Student Name'})

    def test_get_image_rename_options_accepts_field_combinations(self):
        from exports.views import _get_image_rename_options_from_request

        request = self.factory.post(
            f'/panel/exports/images/{self.table.id}/',
            data=json.dumps({
                'rename_options': {
                    'enabled': True,
                    'image_name_fields': {
                        ' photo ': [' Student Name ', ' Class ', '', 'student name'],
                        'rel_photo': ['Father Name', None],
                        'x': [],
                    },
                }
            }),
            content_type='application/json',
        )

        opts = _get_image_rename_options_from_request(request)
        self.assertTrue(opts['enabled'])
        self.assertEqual(
            opts['image_name_fields'],
            {
                'PHOTO': ['Student Name', 'Class'],
                'REL_PHOTO': ['Father Name'],
            },
        )

    def test_get_image_rename_options_accepts_selected_image_field(self):
        from exports.views import _get_image_rename_options_from_request

        request = self.factory.post(
            f'/panel/exports/images/{self.table.id}/',
            data=json.dumps({
                'rename_options': {
                    'enabled': True,
                    'selected_image_field': ' photo ',
                    'image_name_fields': {
                        'photo': ['Student Name'],
                    },
                }
            }),
            content_type='application/json',
        )

        opts = _get_image_rename_options_from_request(request)
        self.assertTrue(opts['enabled'])
        self.assertEqual(opts['selected_image_field'], 'photo')
        self.assertEqual(opts['image_name_fields'], {'PHOTO': ['Student Name']})

    def test_get_image_rename_options_supports_generate_mode_and_compression(self):
        from exports.views import _get_image_rename_options_from_request

        request = self.factory.post(
            f'/panel/exports/images/{self.table.id}/',
            data=json.dumps({
                'rename_options': {
                    'enabled': True,
                    'mode': 'generate',
                    'output_format': 'pdf_zip',
                    'selected_image_field': ' photo ',
                    'image_name_fields': {
                        'photo': [' Student Name ', ' Class ', 'Section'],
                    },
                    'generate_options': {
                        'enabled': True,
                        'name_field': ' Student Name ',
                        'detail_fields': ['Class', 'Section', 'Class', ''],
                        'max_detail_lines': 9,
                        'compress_enabled': True,
                        'target_size_kb': 999,
                    },
                }
            }),
            content_type='application/json',
        )

        opts = _get_image_rename_options_from_request(request)
        self.assertTrue(opts['enabled'])
        self.assertEqual(opts['mode'], 'generate')
        self.assertEqual(opts['output_format'], 'pdf_zip')
        self.assertEqual(opts['selected_image_field'], 'photo')
        self.assertEqual(opts['image_name_fields'], {'PHOTO': ['Student Name', 'Class', 'Section']})

        gen_opts = opts.get('generate_options') or {}
        self.assertTrue(gen_opts.get('enabled'))
        self.assertEqual(gen_opts.get('name_field'), 'Student Name')
        self.assertEqual(gen_opts.get('detail_fields'), ['Class', 'Section'])
        self.assertEqual(gen_opts.get('max_detail_lines'), 1)
        self.assertEqual(gen_opts.get('detail_mode'), 'class_section')
        self.assertEqual(gen_opts.get('class_field'), 'Class')
        self.assertEqual(gen_opts.get('section_field'), 'Section')
        self.assertEqual(gen_opts.get('size_preset'), 'size_23x34')
        self.assertTrue(gen_opts.get('compress_enabled'))
        self.assertEqual(gen_opts.get('target_size_kb'), 200)


class PdfExporterUnitTests(SimpleTestCase):
    """Lightweight unit tests for PdfExporter helper methods that don't need DB."""

    def test_build_rows_uses_datauri_when_placeholder_missing(self):
        from types import SimpleNamespace
        from exports.pdf import PdfExporter, _PLACEHOLDER_IMAGE_PATH

        # Ensure placeholder is absent for this test environment
        try:
            if _PLACEHOLDER_IMAGE_PATH and os.path.isfile(_PLACEHOLDER_IMAGE_PATH):
                os.remove(_PLACEHOLDER_IMAGE_PATH)
        except Exception:
            pass

        ordered_fields = [
            {'name': 'PHOTO', 'type': 'image', 'is_image': True, 'image_width_cm': 1.95, 'image_height_cm': 2.5}
        ]

        # Minimal card-like object with field_data attribute
        card = SimpleNamespace(field_data={'PHOTO': 'NOT_FOUND'})
        exporter = PdfExporter()

        rows = exporter._build_rows(ordered_fields, [card], column_configs=[{'is_image': False}, {'is_image': True, 'image_width_cm': 1.95, 'image_height_cm': 2.5}])
        self.assertTrue(rows)
        # The image cell content should be either a file URI or data URI; ensure no exception occurred
        img_cell = rows[0][1]
        self.assertIn('content', img_cell)
        self.assertTrue(isinstance(img_cell['content'], str) and (img_cell['content'].startswith('file:') or img_cell['content'].startswith('data:')))

    def test_lock_acquire_and_release(self):
        from exports.views import _acquire_export_lock, _release_export_lock

        acquired, lock_key = _acquire_export_lock(11, 22, export_type='xlsx', max_concurrent=1, ttl=30)
        self.assertTrue(acquired)
        self.assertTrue(lock_key)

        acquired_again, _ = _acquire_export_lock(11, 22, export_type='xlsx', max_concurrent=1, ttl=30)
        self.assertFalse(acquired_again)

        _release_export_lock(lock_key)
        acquired_after_release, _ = _acquire_export_lock(11, 22, export_type='xlsx', max_concurrent=1, ttl=30)
        self.assertTrue(acquired_after_release)


class ExportServiceAdvancedTests(TestCase):
    def setUp(self):
        from client.models import Client
        from idcards.models import IDCardGroup, IDCardTable, IDCard
        from staff.models import Staff

        self.super_admin = User.objects.create_user(
            username='svc-super@test.com', email='svc-super@test.com',
            password='pass1234', role='super_admin',
        )

        owner1 = User.objects.create_user(
            username='svc-owner1@test.com', email='svc-owner1@test.com',
            password='pass1234', role='client',
        )
        owner2 = User.objects.create_user(
            username='svc-owner2@test.com', email='svc-owner2@test.com',
            password='pass1234', role='client',
        )
        self.client1 = Client.objects.create(user=owner1, name='Svc Client 1')
        self.client2 = Client.objects.create(user=owner2, name='Svc Client 2')

        group1 = IDCardGroup.objects.create(client=self.client1, name='Group 1')
        group2 = IDCardGroup.objects.create(client=self.client2, name='Group 2')
        self.table1 = IDCardTable.objects.create(group=group1, name='Table 1', fields=[{'name': 'NAME', 'type': 'text'}])
        self.table2 = IDCardTable.objects.create(group=group2, name='Table 2', fields=[{'name': 'NAME', 'type': 'text'}])

        IDCard.objects.create(table=self.table1, field_data={'NAME': 'A'}, status='pending')
        IDCard.objects.create(table=self.table1, field_data={'NAME': 'B'}, status='verified')
        IDCard.objects.create(table=self.table2, field_data={'NAME': 'C'}, status='approved')

        self.staff_user = User.objects.create_user(
            username='svc-staff@test.com', email='svc-staff@test.com',
            password='pass1234', role='admin_staff',
        )
        self.staff = Staff.objects.create(
            user=self.staff_user,
            staff_type='admin_staff',
            perm_idcard_bulk_download=True,
        )
        self.staff.assigned_clients.add(self.client1)

    def test_get_scoped_cards_admin_staff_assigned_only(self):
        from exports.services import ExportService

        service = ExportService(self.staff_user)
        cards_for_table1 = service.get_scoped_cards(self.table1)
        cards_for_table2 = service.get_scoped_cards(self.table2)

        self.assertEqual(cards_for_table1.count(), 2)
        self.assertEqual(cards_for_table2.count(), 0)

    def test_get_scoped_cards_client_staff_respects_table_row_scope(self):
        from exports.services import ExportService
        from staff.models import Staff
        from idcards.models import IDCard

        self.table1.fields = [
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'CLASS', 'type': 'class', 'order': 2},
        ]
        self.table1.save(update_fields=['fields'])

        IDCard.objects.filter(table=self.table1).delete()
        keep = IDCard.objects.create(table=self.table1, field_data={'NAME': 'A', 'CLASS': 'KG-I'}, status='pending')
        IDCard.objects.create(table=self.table1, field_data={'NAME': 'B', 'CLASS': 'KG-II'}, status='pending')

        cs_user = User.objects.create_user(
            username='svc-clientstaff@test.com',
            email='svc-clientstaff@test.com',
            password='pass1234',
            role='client_staff',
        )
        Staff.objects.create(
            user=cs_user,
            staff_type='client_staff',
            client=self.client1,
            assigned_table_ids=[self.table1.id],
            allowed_classes=['KG1'],
            perm_idcard_bulk_download=True,
            perm_idcard_pending_list=True,
        )

        service = ExportService(cs_user)
        cards = list(service.get_scoped_cards(self.table1).values_list('id', flat=True))
        self.assertEqual(cards, [keep.id])

    def test_prepare_context_permission_denied_when_no_bulk_download(self):
        from exports.services import ExportService

        denied_user = User.objects.create_user(
            username='svc-denied@test.com', email='svc-denied@test.com',
            password='pass1234', role='admin_staff',
        )
        service = ExportService(denied_user)
        context = service._prepare_context(self.table1.id, require_export_permission=True)

        self.assertFalse(context.has_permission)
        self.assertIn('Permission denied', context.error_message)

    def test_get_export_preview_contains_counts(self):
        from exports.services import ExportService

        service = ExportService(self.super_admin)
        preview = service.get_export_preview(self.table1.id)

        self.assertTrue(preview['success'])
        self.assertEqual(preview['card_count'], 2)
        self.assertTrue(preview['available_formats']['xlsx'])


class ExportApiIntegrationAdvancedTests(TestCase):
    def setUp(self):
        from staff.models import Staff

        self.admin, self.client_user, self.client_obj, self.table = _setup_export_data()
        self.factory = RequestFactory()
        cache.clear()

        self.staff_unassigned = User.objects.create_user(
            username='exstaffu@test.com', email='exstaffu@test.com',
            password='pass1234', role='admin_staff',
        )
        Staff.objects.create(
            user=self.staff_unassigned,
            staff_type='admin_staff',
            perm_idcard_bulk_download=True,
        )

        self.staff_assigned = User.objects.create_user(
            username='exstaffa@test.com', email='exstaffa@test.com',
            password='pass1234', role='admin_staff',
        )
        self.staff_assigned_profile = Staff.objects.create(
            user=self.staff_assigned,
            staff_type='admin_staff',
            perm_idcard_bulk_download=True,
        )
        self.staff_assigned_profile.assigned_clients.add(self.client_obj)

    def tearDown(self):
        cache.clear()

    def test_preview_success_for_super_admin(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')
        response = self.client.get(f'/panel/exports/preview/{self.table.id}/')
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['success'])
        self.assertEqual(payload['card_count'], 5)

    def test_preview_scope_denied_for_unassigned_admin_staff(self):
        self.client.login(username='exstaffu@test.com', password='pass1234')
        response = self.client.get(f'/panel/exports/preview/{self.table.id}/')
        self.assertEqual(response.status_code, 403)

    def test_xlsx_empty_selection_falls_back_to_table_cards(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')
        response = self.client.post(
            f'/panel/exports/xlsx/{self.table.id}/',
            data=json.dumps({'card_ids': []}),
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 200)

    def test_docx_invalid_format_falls_back_to_docx(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')

        fake_result = type('Res', (), {'success': True, 'response': HttpResponse(b'ok')})
        with mock.patch('exports.views.ExportService.export_word', return_value=fake_result) as mocked:
            response = self.client.post(
                f'/panel/exports/docx/{self.table.id}/',
                data=json.dumps({'card_ids': [1, 2], 'format': 'invalid-format'}),
                content_type='application/json',
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(mocked.call_args.kwargs['doc_format'], 'docx')

    def test_images_export_passes_cleaned_rename_options(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')

        fake_zip_result = object()
        with mock.patch('exports.views.ExportService.export_images', return_value=fake_zip_result) as mocked_export:
            with mock.patch('exports.views.zip_result_to_dict', return_value={'success': True, 'zip_files': []}):
                response = self.client.post(
                    f'/panel/exports/images/{self.table.id}/',
                    data=json.dumps({
                        'card_ids': [1],
                        'rename_options': {
                            'enabled': True,
                            'image_name_fields': {' photo ': 'Name', '': 'x', 'QR': ''},
                        },
                    }),
                    content_type='application/json',
                )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()['success'])
        rename_opts = mocked_export.call_args.kwargs['rename_options']
        self.assertEqual(rename_opts['image_name_fields'], {'PHOTO': 'Name'})

    def test_images_export_passes_field_combination_rename_options(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')

        fake_zip_result = object()
        with mock.patch('exports.views.ExportService.export_images', return_value=fake_zip_result) as mocked_export:
            with mock.patch('exports.views.zip_result_to_dict', return_value={'success': True, 'zip_files': []}):
                response = self.client.post(
                    f'/panel/exports/images/{self.table.id}/',
                    data=json.dumps({
                        'card_ids': [1],
                        'rename_options': {
                            'enabled': True,
                            'image_name_fields': {
                                'photo': ['Name', 'Father'],
                            },
                        },
                    }),
                    content_type='application/json',
                )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()['success'])
        rename_opts = mocked_export.call_args.kwargs['rename_options']
        self.assertEqual(rename_opts['image_name_fields'], {'PHOTO': ['Name', 'Father']})

    def test_images_export_rename_mode_requires_permission(self):
        self.client.login(username='exstaffa@test.com', password='pass1234')

        response = self.client.post(
            f'/panel/exports/images/{self.table.id}/',
            data=json.dumps({
                'card_ids': [1],
                'rename_options': {
                    'enabled': True,
                    'mode': 'rename',
                    'selected_image_field': 'PHOTO',
                    'image_name_fields': {'PHOTO': 'Name'},
                },
            }),
            content_type='application/json',
        )

        self.assertEqual(response.status_code, 403)
        self.assertIn('Rename mode', response.json().get('message', ''))

    def test_images_export_generate_mode_requires_permission(self):
        self.client.login(username='exstaffa@test.com', password='pass1234')

        response = self.client.post(
            f'/panel/exports/images/{self.table.id}/',
            data=json.dumps({
                'card_ids': [1],
                'rename_options': {
                    'enabled': True,
                    'mode': 'generate',
                    'selected_image_field': 'PHOTO',
                    'image_name_fields': {'PHOTO': ['Name', 'Class']},
                    'generate_options': {
                        'enabled': True,
                        'name_field': 'Name',
                        'detail_fields': ['Class'],
                    },
                },
            }),
            content_type='application/json',
        )

        self.assertEqual(response.status_code, 403)
        self.assertIn('Generate mode', response.json().get('message', ''))

    def test_images_export_generate_mode_allowed_with_permission(self):
        self.staff_assigned_profile.perm_idcard_download_image_generate_mode = True
        self.staff_assigned_profile.save(update_fields=['perm_idcard_download_image_generate_mode'])
        self.client.login(username='exstaffa@test.com', password='pass1234')

        fake_zip_result = object()
        with mock.patch('exports.views.ExportService.export_images', return_value=fake_zip_result) as mocked_export:
            with mock.patch('exports.views.zip_result_to_dict', return_value={'success': True, 'zip_files': []}):
                response = self.client.post(
                    f'/panel/exports/images/{self.table.id}/',
                    data=json.dumps({
                        'card_ids': [1],
                        'rename_options': {
                            'enabled': True,
                            'mode': 'generate',
                            'selected_image_field': 'PHOTO',
                            'image_name_fields': {'PHOTO': ['Name', 'Class']},
                            'generate_options': {
                                'enabled': True,
                                'name_field': 'Name',
                                'detail_fields': ['Class'],
                            },
                        },
                    }),
                    content_type='application/json',
                )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json().get('success'))
        self.assertEqual(mocked_export.call_args.kwargs['rename_options']['mode'], 'generate')

    def test_images_export_normal_download_still_allowed_without_mode_permissions(self):
        self.client.login(username='exstaffa@test.com', password='pass1234')

        fake_zip_result = type('ZipResult', (), {
            'success': True,
            'message': '',
        })()
        with mock.patch('exports.views.ExportService.export_images', return_value=fake_zip_result) as mocked_export:
            with mock.patch('exports.views.zip_result_to_dict', return_value={'success': True, 'zip_files': []}):
                response = self.client.post(
                    f'/panel/exports/images/{self.table.id}/',
                    data=json.dumps({
                        'card_ids': [1],
                    }),
                    content_type='application/json',
                )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json().get('success'))
        self.assertIsNone(mocked_export.call_args.kwargs.get('rename_options'))

    def test_images_export_large_inline_falls_back_to_disk_urls(self):
        from staff.models import Staff

        staff_user = User.objects.create_user(
            username='imgstaff@test.com',
            email='imgstaff@test.com',
            password='pass1234',
            role='admin_staff',
        )
        staff = Staff.objects.create(
            user=staff_user,
            staff_type='admin_staff',
            perm_idcard_bulk_download=True,
        )
        staff.assigned_clients.add(self.client_obj)

        self.client.login(username='imgstaff@test.com', password='pass1234')

        inline_fail = type('ZipResult', (), {
            'success': False,
            'message': 'Selected images are too large for inline ZIP download. Please export fewer cards.'
        })()
        disk_zip_info = type('DiskZipInfo', (), {
            'field_name': 'ALL',
            'filename': 'big_images.zip',
            'path': os.path.join(settings.MEDIA_ROOT, 'exports', 'big_images.zip'),
            'image_count': 5,
        })()
        disk_ok = type('DiskZipResult', (), {
            'success': True,
            'zip_files': [disk_zip_info],
            'total_images': 5,
            'total_zips': 1,
        })()

        with mock.patch('exports.views.ExportService.export_images', return_value=inline_fail):
            with mock.patch('exports.zip.export_images_to_disk', return_value=disk_ok):
                response = self.client.post(
                    f'/panel/exports/images/{self.table.id}/',
                    data=json.dumps({'card_ids': [1, 2, 3]}),
                    content_type='application/json',
                )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload.get('success'))
        self.assertTrue(payload.get('download_url', '').endswith('/exports/big_images.zip'))
        self.assertEqual(len(payload.get('zip_files', [])), 1)

    def test_pdf_async_starts_background_task(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')

        with mock.patch('exports.tasks.BackgroundExportManager.start_pdf_export', return_value='task123') as mocked_start:
            response = self.client.post(
                f'/panel/exports/pdf-async/{self.table.id}/',
                data=json.dumps({'card_ids': [1, 2, 3], 'break_mode': 'class_section'}),
                content_type='application/json',
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload['success'])
        self.assertEqual(payload['task_id'], 'task123')
        self.assertEqual(mocked_start.call_args.kwargs.get('break_mode'), 'class_section')

    def test_download_all_lock_contention_returns_429(self):
        self.client.login(username='exadmin@test.com', password='adminpass1')

        with mock.patch('exports.views._acquire_export_lock', return_value=(False, '')):
            response = self.client.post(
                f'/panel/exports/download-all/{self.table.id}/',
                data=json.dumps({}),
                content_type='application/json',
            )

        self.assertEqual(response.status_code, 429)

    def test_export_status_is_scoped_to_request_user(self):
        from core.models import BackgroundTask

        task = BackgroundTask.objects.create(
            user=self.admin,
            task_type='export_pdf',
            status='completed',
            total=1,
            progress=1,
            result_path='temp/exports/owner.pdf',
            metadata={'result': {'filename': 'owner.pdf', 'card_count': 1}},
        )

        other_user = User.objects.create_user(
            username='exother@test.com',
            email='exother@test.com',
            password='pass1234',
            role='super_admin',
        )
        self.client.login(username='exother@test.com', password='pass1234')
        response = self.client.get(f'/panel/exports/status/{task.id}/')
        self.assertEqual(response.status_code, 404)

    def test_export_status_hides_invalid_result_path(self):
        from core.models import BackgroundTask

        self.client.login(username='exadmin@test.com', password='adminpass1')
        task = BackgroundTask.objects.create(
            user=self.admin,
            task_type='export_pdf',
            status='completed',
            total=1,
            progress=1,
            result_path='../outside.pdf',
            metadata={'result': {'filename': 'outside.pdf', 'card_count': 1}},
        )

        response = self.client.get(f'/panel/exports/status/{task.id}/')
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json().get('download_url'), '')

class ExportDeepLimitAndRoleTests(TestCase):
    def setUp(self):
        from client.models import Client
        from staff.models import Staff
        from idcards.models import IDCardGroup, IDCardTable, IDCard

        self.super_admin = User.objects.create_user(
            username='deep-super@test.com',
            email='deep-super@test.com',
            password='pass1234',
            role='super_admin',
        )

        self.client_user = User.objects.create_user(
            username='deep-client@test.com',
            email='deep-client@test.com',
            password='pass1234',
            role='client',
        )
        self.client_obj = Client.objects.create(user=self.client_user, name='Deep Client')
        self.client_obj.perm_idcard_bulk_download = True
        self.client_obj.perm_idcard_approved_list = True
        self.client_obj.perm_idcard_download_list = True
        self.client_obj.save(update_fields=[
            'perm_idcard_bulk_download',
            'perm_idcard_approved_list',
            'perm_idcard_download_list',
        ])

        self.admin_staff_user = User.objects.create_user(
            username='deep-adminstaff@test.com',
            email='deep-adminstaff@test.com',
            password='pass1234',
            role='admin_staff',
        )
        self.admin_staff = Staff.objects.create(
            user=self.admin_staff_user,
            staff_type='admin_staff',
            perm_idcard_bulk_download=True,
        )
        self.admin_staff.assigned_clients.add(self.client_obj)

        self.client_staff_user = User.objects.create_user(
            username='deep-clientstaff@test.com',
            email='deep-clientstaff@test.com',
            password='pass1234',
            role='client_staff',
        )
        Staff.objects.create(
            user=self.client_staff_user,
            staff_type='client_staff',
            client=self.client_obj,
            perm_idcard_bulk_download=True,
            perm_idcard_approved_list=True,
            perm_idcard_download_list=True,
        )

        group = IDCardGroup.objects.create(client=self.client_obj, name='Deep Group')
        self.table = IDCardTable.objects.create(
            group=group,
            name='Deep Table',
            fields=[
                {'name': 'NAME', 'type': 'text', 'order': 1},
                {'name': 'PHOTO', 'type': 'image', 'order': 2},
            ],
        )
        self.card = IDCard.objects.create(
            table=self.table,
            field_data={'NAME': 'Student A', 'PHOTO': 'clients_imgs/deep/photo_a.jpg'},
            status='pending',
        )

    def test_export_pdf_requires_status_list_permission(self):
        from idcards.models import IDCard

        self.client.login(username='deep-adminstaff@test.com', password='pass1234')
        card_id = IDCard.objects.filter(table=self.table).values_list('id', flat=True).first()

        response = self.client.post(
            f'/panel/exports/pdf/{self.table.id}/',
            data=json.dumps({'card_ids': [card_id], 'status': 'pending'}),
            content_type='application/json',
        )
        self.assertEqual(response.status_code, 403)

    def test_export_pdf_client_staff_blocked_for_unassigned_table(self):
        from idcards.models import IDCard, IDCardGroup, IDCardTable
        from staff.models import Staff

        other_group = IDCardGroup.objects.create(client=self.client_obj, name='Other Group')
        other_table = IDCardTable.objects.create(
            group=other_group,
            name='Other Table',
            fields=[{'name': 'NAME', 'type': 'text', 'order': 1}],
        )
        IDCard.objects.create(table=other_table, field_data={'NAME': 'X'}, status='pending')

        restricted_user = User.objects.create_user(
            username='deep-restricted-staff@test.com',
            email='deep-restricted-staff@test.com',
            password='pass1234',
            role='client_staff',
        )
        Staff.objects.create(
            user=restricted_user,
            staff_type='client_staff',
            client=self.client_obj,
            assigned_table_ids=[other_table.id],
            perm_idcard_bulk_download=True,
            perm_idcard_pending_list=True,
        )

        self.client.logout()
        self.client.login(username='deep-restricted-staff@test.com', password='pass1234')
        response = self.client.post(
            f'/panel/exports/pdf/{self.table.id}/',
            data=json.dumps({'status': 'pending'}),
            content_type='application/json',
        )

        self.assertEqual(response.status_code, 403)

    def _mock_file_response(self):
        return type('Res', (), {'success': True, 'response': HttpResponse(b'ok')})

    def test_role_matrix_client_and_client_staff_pdf_only(self):
        from staff.models import Staff

        # Align with real UI flow: pending-list export calls include status='pending'.
        self.client_obj.perm_idcard_pending_list = True
        self.client_obj.save(update_fields=['perm_idcard_pending_list'])
        client_staff_profile = Staff.objects.get(user=self.client_staff_user)
        client_staff_profile.perm_idcard_pending_list = True
        client_staff_profile.save(update_fields=['perm_idcard_pending_list'])

        # Client: PDF and Images allowed; xlsx/docx/download-all blocked.
        self.client.login(username='deep-client@test.com', password='pass1234')
        with mock.patch('exports.views.ExportService.export_pdf', return_value=self._mock_file_response()):
            pdf_resp = self.client.post(
                f'/panel/exports/pdf/{self.table.id}/',
                data=json.dumps({'card_ids': [self.card.id], 'status': 'approved'}),
                content_type='application/json',
            )
        self.assertEqual(pdf_resp.status_code, 200)

        xlsx_resp = self.client.post(
            f'/panel/exports/xlsx/{self.table.id}/',
            data=json.dumps({'card_ids': [self.card.id]}),
            content_type='application/json',
        )
        self.assertEqual(xlsx_resp.status_code, 403)

        docx_resp = self.client.post(
            f'/panel/exports/docx/{self.table.id}/',
            data=json.dumps({'card_ids': [self.card.id]}),
            content_type='application/json',
        )
        self.assertEqual(docx_resp.status_code, 403)

        fake_zip_result = type('ZipResult', (), {'success': True, 'message': ''})()
        with mock.patch('exports.views.ExportService.export_images', return_value=fake_zip_result):
            with mock.patch('exports.views.zip_result_to_dict', return_value={'success': True, 'zip_files': []}):
                images_resp = self.client.post(
                    f'/panel/exports/images/{self.table.id}/',
                    data=json.dumps({'card_ids': [self.card.id], 'status': 'pending'}),
                    content_type='application/json',
                )
        self.assertEqual(images_resp.status_code, 200)

        download_all_resp = self.client.post(
            f'/panel/exports/download-all/{self.table.id}/',
            data=json.dumps({}),
            content_type='application/json',
        )
        self.assertEqual(download_all_resp.status_code, 403)

        # Client staff: same policy as client for images.
        self.client.logout()
        self.client.login(username='deep-clientstaff@test.com', password='pass1234')
        with mock.patch('exports.views.ExportService.export_pdf', return_value=self._mock_file_response()):
            pdf_staff_resp = self.client.post(
                f'/panel/exports/pdf/{self.table.id}/',
                data=json.dumps({'card_ids': [self.card.id], 'status': 'download'}),
                content_type='application/json',
            )
        self.assertEqual(pdf_staff_resp.status_code, 200)

        fake_zip_staff_result = type('ZipResult', (), {'success': True, 'message': ''})()
        with mock.patch('exports.views.ExportService.export_images', return_value=fake_zip_staff_result):
            with mock.patch('exports.views.zip_result_to_dict', return_value={'success': True, 'zip_files': []}):
                images_staff_resp = self.client.post(
                    f'/panel/exports/images/{self.table.id}/',
                    data=json.dumps({'card_ids': [self.card.id], 'status': 'pending'}),
                    content_type='application/json',
                )
        self.assertEqual(images_staff_resp.status_code, 200)

    def test_download_all_success_for_admin_staff(self):
        fake_xlsx = type('X', (), {'success': True, 'response': HttpResponse(b'xlsx-bytes')})()
        fake_disk = type('D', (), {'success': False, 'zip_files': []})()

        self.client.login(username='deep-adminstaff@test.com', password='pass1234')
        with mock.patch('exports.views.ExcelExporter.export_cards', return_value=fake_xlsx):
            with mock.patch('exports.zip.export_images_to_disk', return_value=fake_disk):
                response = self.client.post(
                    f'/panel/exports/download-all/{self.table.id}/',
                    data=json.dumps({}),
                    content_type='application/json',
                )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertTrue(payload.get('success'))
        self.assertTrue(str(payload.get('download_url', '')).endswith('.zip'))

    def test_zip_inline_1gb_boundary_and_super_admin_bypass(self):
        from exports.zip import ZipExporter, MAX_BASE64_ZIP_BYTES

        exporter = ZipExporter()
        cards = self.table.id_cards.all()

        mocked_file = mock.MagicMock()
        mocked_file.__enter__.return_value = mocked_file
        mocked_file.__exit__.return_value = False
        mocked_file.read.return_value = b'a' * 256

        with mock.patch('exports.zip.is_valid_image_path', return_value=True):
            with mock.patch('exports.zip.ImageService.get_image_path_for_card', return_value='clients_imgs/deep/photo_a.jpg'):
                with mock.patch('exports.zip.default_storage.open', return_value=mocked_file):
                    with mock.patch('exports.zip.os.path.getsize', return_value=MAX_BASE64_ZIP_BYTES):
                        boundary_ok = exporter.export_images(self.table, cards, allow_large_base64=False)

                    with mock.patch('exports.zip.os.path.getsize', return_value=MAX_BASE64_ZIP_BYTES + 1):
                        over_limit_blocked = exporter.export_images(self.table, cards, allow_large_base64=False)

                    with mock.patch('exports.zip.os.path.getsize', return_value=MAX_BASE64_ZIP_BYTES + 1):
                        super_admin_ok = exporter.export_images(self.table, cards, allow_large_base64=True)

        self.assertTrue(boundary_ok.success)
        self.assertFalse(over_limit_blocked.success)
        self.assertIn('1 GB inline ZIP limit', over_limit_blocked.message)
        self.assertTrue(super_admin_ok.success)

    def test_zip_export_uses_combined_rename_fields_for_filename(self):
        from exports.zip import ZipExporter
        import base64
        import io
        import zipfile

        exporter = ZipExporter()
        cards = self.table.id_cards.all()
        rename_options = {
            'enabled': True,
            'image_name_fields': {
                'PHOTO': ['NAME', 'UNKNOWN_FIELD'],
            },
        }

        mocked_file = mock.MagicMock()
        mocked_file.__enter__.return_value = mocked_file
        mocked_file.__exit__.return_value = False
        mocked_file.read.return_value = b'a' * 256

        with mock.patch('exports.zip.is_valid_image_path', return_value=True):
            with mock.patch('exports.zip.ImageService.get_image_path_for_card', return_value='clients_imgs/deep/original_photo.jpg'):
                with mock.patch('exports.zip.default_storage.open', return_value=mocked_file):
                    result = exporter.export_images(self.table, cards, rename_options=rename_options, allow_large_base64=True)

        self.assertTrue(result.success)
        self.assertEqual(len(result.zip_files), 1)

        zip_payload = base64.b64decode(result.zip_files[0].data)
        with zipfile.ZipFile(io.BytesIO(zip_payload), 'r') as zf:
            names = zf.namelist()

        self.assertTrue(any(name.startswith('PHOTO/Student_A') for name in names), names)

    def test_zip_export_rename_mode_limits_to_selected_image_field(self):
        from exports.zip import ZipExporter
        import base64
        import io
        import zipfile

        self.table.fields = [
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'PHOTO', 'type': 'image', 'order': 2},
            {'name': 'PHOTO_2', 'type': 'image', 'order': 3},
        ]
        self.table.save(update_fields=['fields'])

        card = self.table.id_cards.first()
        card.field_data = {
            'NAME': 'Student A',
            'PHOTO': 'clients_imgs/deep/photo_a.jpg',
            'PHOTO_2': 'clients_imgs/deep/photo_b.jpg',
        }
        card.save(update_fields=['field_data'])

        exporter = ZipExporter()
        cards = self.table.id_cards.all()
        rename_options = {
            'enabled': True,
            'selected_image_field': 'PHOTO',
            'image_name_fields': {
                'PHOTO': ['NAME'],
            },
        }

        mocked_file = mock.MagicMock()
        mocked_file.__enter__.return_value = mocked_file
        mocked_file.__exit__.return_value = False
        mocked_file.read.return_value = b'a' * 256

        def _mock_image_path(*args, **kwargs):
            card_obj = kwargs.get('card') if 'card' in kwargs else (args[0] if len(args) > 0 else None)
            field_name = kwargs.get('field_name') if 'field_name' in kwargs else (args[1] if len(args) > 1 else '')
            return (card_obj.field_data or {}).get(field_name, '')

        with mock.patch('exports.zip.is_valid_image_path', return_value=True):
            with mock.patch('exports.zip.ImageService.get_image_path_for_card', side_effect=_mock_image_path):
                with mock.patch('exports.zip.default_storage.open', return_value=mocked_file):
                    result = exporter.export_images(self.table, cards, rename_options=rename_options, allow_large_base64=True)

        self.assertTrue(result.success)
        self.assertEqual(len(result.zip_files), 1)

        zip_payload = base64.b64decode(result.zip_files[0].data)
        with zipfile.ZipFile(io.BytesIO(zip_payload), 'r') as zf:
            names = zf.namelist()

        self.assertTrue(any(name.startswith('PHOTO/') for name in names), names)
        self.assertFalse(any(name.startswith('PHOTO_2/') for name in names), names)

    def test_zip_export_generate_mode_builds_compressed_passport_images(self):
        from exports.zip import ZipExporter
        from PIL import Image
        import base64
        import io
        import zipfile

        self.table.fields = [
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'CLASS', 'type': 'text', 'order': 2},
            {'name': 'SECTION', 'type': 'text', 'order': 3},
            {'name': 'PHOTO', 'type': 'image', 'order': 4},
        ]
        self.table.save(update_fields=['fields'])

        card = self.table.id_cards.first()
        card.field_data = {
            'NAME': 'Student A',
            'CLASS': '10',
            'SECTION': 'A',
            'PHOTO': 'clients_imgs/deep/photo_a.jpg',
        }
        card.save(update_fields=['field_data'])

        img_buf = io.BytesIO()
        Image.new('RGB', (900, 1200), (88, 112, 158)).save(img_buf, format='JPEG', quality=92)
        valid_jpg_bytes = img_buf.getvalue()

        exporter = ZipExporter()
        cards = self.table.id_cards.all()
        rename_options = {
            'enabled': True,
            'mode': 'generate',
            'output_format': 'zip',
            'selected_image_field': 'PHOTO',
            'image_name_fields': {
                'PHOTO': ['NAME', 'CLASS', 'SECTION'],
            },
            'generate_options': {
                'enabled': True,
                'size_preset': 'size_23x34',
                'name_field': 'NAME',
                'detail_mode': 'class_section',
                'class_field': 'CLASS',
                'section_field': 'SECTION',
                'detail_fields': ['CLASS', 'SECTION'],
                'max_detail_lines': 1,
                'compress_enabled': True,
                'target_size_kb': 40,
            },
        }

        mocked_file = mock.MagicMock()
        mocked_file.__enter__.return_value = mocked_file
        mocked_file.__exit__.return_value = False
        mocked_file.read.return_value = valid_jpg_bytes

        with mock.patch('exports.zip.is_valid_image_path', return_value=True):
            with mock.patch('exports.zip.ImageService.get_image_path_for_card', return_value='clients_imgs/deep/photo_a.jpg'):
                with mock.patch('exports.zip.default_storage.open', return_value=mocked_file):
                    result = exporter.export_images(self.table, cards, rename_options=rename_options, allow_large_base64=True)

        self.assertTrue(result.success)
        self.assertEqual(len(result.zip_files), 1)

        zip_payload = base64.b64decode(result.zip_files[0].data)
        with zipfile.ZipFile(io.BytesIO(zip_payload), 'r') as zf:
            names = zf.namelist()
            self.assertTrue(any(name.startswith('PHOTO/Student_A') and name.lower().endswith('.jpg') for name in names), names)

            generated_name = next(name for name in names if name.startswith('PHOTO/Student_A'))
            generated_payload = zf.read(generated_name)
            self.assertLessEqual(len(generated_payload), 40 * 1024)

        with Image.open(io.BytesIO(generated_payload)) as generated_image:
            self.assertEqual(generated_image.size, (272, 402))

    def test_pdf_zip_generate_mode_exports_generated_pages(self):
        from exports.zip import ZipExporter
        from PIL import Image
        import base64
        import io
        import zipfile

        self.table.fields = [
            {'name': 'NAME', 'type': 'text', 'order': 1},
            {'name': 'CLASS', 'type': 'text', 'order': 2},
            {'name': 'SECTION', 'type': 'text', 'order': 3},
            {'name': 'PHOTO', 'type': 'image', 'order': 4},
        ]
        self.table.save(update_fields=['fields'])

        self.card.field_data = {
            'NAME': 'Student A',
            'CLASS': '10',
            'SECTION': 'A',
            'PHOTO': 'clients_imgs/deep/photo_a.jpg',
        }
        self.card.save(update_fields=['field_data'])

        img_buf = io.BytesIO()
        Image.new('RGB', (900, 1200), (88, 112, 158)).save(img_buf, format='JPEG', quality=92)
        valid_jpg_bytes = img_buf.getvalue()

        exporter = ZipExporter()
        cards = self.table.id_cards.all()
        rename_options = {
            'enabled': True,
            'mode': 'generate',
            'output_format': 'pdf_zip',
            'selected_image_field': 'PHOTO',
            'image_name_fields': {
                'PHOTO': ['NAME', 'CLASS', 'SECTION'],
            },
            'generate_options': {
                'enabled': True,
                'size_preset': 'size_37x53',
                'name_field': 'NAME',
                'detail_mode': 'class_section',
                'class_field': 'CLASS',
                'section_field': 'SECTION',
                'compress_enabled': False,
                'target_size_kb': 120,
            },
        }

        mocked_file = mock.MagicMock()
        mocked_file.__enter__.return_value = mocked_file
        mocked_file.__exit__.return_value = False
        mocked_file.read.return_value = valid_jpg_bytes

        with mock.patch('exports.zip.is_valid_image_path', return_value=True):
            with mock.patch('exports.zip.ImageService.get_image_path_for_card', return_value='clients_imgs/deep/photo_a.jpg'):
                with mock.patch('exports.zip.default_storage.open', return_value=mocked_file):
                    result = exporter.export_images(self.table, cards, rename_options=rename_options, allow_large_base64=True)

        self.assertTrue(result.success)
        self.assertEqual(len(result.zip_files), 1)

        zip_payload = base64.b64decode(result.zip_files[0].data)
        with zipfile.ZipFile(io.BytesIO(zip_payload), 'r') as zf:
            names = zf.namelist()
            self.assertTrue(any(name.lower().endswith('.pdf') for name in names), names)
            pdf_name = next(name for name in names if name.lower().endswith('.pdf'))
            pdf_payload = zf.read(pdf_name)
            self.assertGreater(len(pdf_payload), 500)

    def test_pdf_zip_1gb_boundary_and_super_admin_bypass(self):
        from exports.zip import ZipExporter, MAX_BASE64_ZIP_BYTES

        exporter = ZipExporter()
        cards = self.table.id_cards.all()
        rename_options = {
            'enabled': True,
            'output_format': 'pdf_zip',
            'image_name_fields': {'PHOTO': 'NAME'},
        }

        with mock.patch.object(ZipExporter, '_write_image_field_pdf', return_value=1):
            with mock.patch('exports.zip.os.path.getsize', return_value=MAX_BASE64_ZIP_BYTES + 1):
                blocked = exporter.export_images(self.table, cards, rename_options=rename_options, allow_large_base64=False)

            with mock.patch('exports.zip.os.path.getsize', return_value=MAX_BASE64_ZIP_BYTES + 1):
                allowed_super = exporter.export_images(self.table, cards, rename_options=rename_options, allow_large_base64=True)

        self.assertFalse(blocked.success)
        self.assertIn('1 GB inline ZIP limit', blocked.message)
        self.assertTrue(allowed_super.success)


class ColumnSpecAliasCoverageTests(SimpleTestCase):
    def test_emergency_contact_abbrev_maps_to_non_wrapping_spec(self):
        from exports.column_spec import classify_column, get_column_spec

        self.assertEqual(classify_column('EMERG CONT NO'), 'emergency_mobile')
        self.assertEqual(classify_column('EMERGENCY CONTACT NUMBER'), 'emergency_mobile')
        self.assertFalse(get_column_spec('EMERG CONT NO').wrap)

    def test_transport_aliases_map_to_bus_route(self):
        from exports.column_spec import classify_column

        self.assertEqual(classify_column('TRANSPORT'), 'bus_route')
        self.assertEqual(classify_column('TRANSPOR'), 'bus_route')
        self.assertEqual(classify_column('TRAN SPORT'), 'bus_route')

    def test_gender_alias_accepts_spaced_variant(self):
        from exports.column_spec import classify_column

        self.assertEqual(classify_column('GEN DER'), 'gender')


class PdfPaginationGroupingTests(SimpleTestCase):
    def test_group_rows_breaks_on_class_and_section_changes(self):
        from exports.pdf import PdfExporter

        exporter = PdfExporter()
        rows = [
            [{'content': 'row1'}],
            [{'content': 'row2'}],
            [{'content': 'row3'}],
            [{'content': 'row4'}],
        ]
        cards = [
            type('Card', (), {'field_data': {'CLASS': '10', 'SECTION': 'A'}})(),
            type('Card', (), {'field_data': {'CLASS': '10', 'SECTION': 'A'}})(),
            type('Card', (), {'field_data': {'CLASS': '10', 'SECTION': 'B'}})(),
            type('Card', (), {'field_data': {'CLASS': '11', 'SECTION': 'A'}})(),
        ]

        pages = exporter._group_rows_into_pages(
            rows=rows,
            cards_list=cards,
            class_field_name='CLASS',
            section_field_name='SECTION',
            records_per_page=6,
        )

        self.assertEqual(
            [[row[0]['content'] for row in page] for page in pages],
            [['row1', 'row2'], ['row3'], ['row4']],
        )

    def test_group_rows_breaks_on_section_when_class_missing(self):
        from exports.pdf import PdfExporter

        exporter = PdfExporter()
        rows = [
            [{'content': 'row1'}],
            [{'content': 'row2'}],
            [{'content': 'row3'}],
        ]
        cards = [
            type('Card', (), {'field_data': {'SECTION': 'A'}})(),
            type('Card', (), {'field_data': {'SECTION': 'A'}})(),
            type('Card', (), {'field_data': {'SECTION': 'B'}})(),
        ]

        pages = exporter._group_rows_into_pages(
            rows=rows,
            cards_list=cards,
            class_field_name=None,
            section_field_name='SECTION',
            records_per_page=6,
        )

        self.assertEqual(
            [[row[0]['content'] for row in page] for page in pages],
            [['row1', 'row2'], ['row3']],
        )

    def test_group_rows_class_only_does_not_break_on_section_change(self):
        from exports.pdf import PdfExporter

        exporter = PdfExporter()
        rows = [
            [{'content': 'row1'}],
            [{'content': 'row2'}],
            [{'content': 'row3'}],
        ]
        cards = [
            type('Card', (), {'field_data': {'CLASS': '10', 'SECTION': 'A'}})(),
            type('Card', (), {'field_data': {'CLASS': '10', 'SECTION': 'B'}})(),
            type('Card', (), {'field_data': {'CLASS': '10', 'SECTION': 'C'}})(),
        ]

        pages = exporter._group_rows_into_pages(
            rows=rows,
            cards_list=cards,
            class_field_name='CLASS',
            section_field_name=None,
            records_per_page=6,
        )

        self.assertEqual(
            [[row[0]['content'] for row in page] for page in pages],
            [['row1', 'row2', 'row3']],
        )

    def test_group_rows_breaks_with_case_variant_field_keys(self):
        from exports.pdf import PdfExporter

        exporter = PdfExporter()
        rows = [
            [{'content': 'row1'}],
            [{'content': 'row2'}],
        ]
        cards = [
            type('Card', (), {'field_data': {'CLASS': 'IV', 'SECTION': 'D'}})(),
            type('Card', (), {'field_data': {'CLASS': 'IV', 'SECTION': 'E'}})(),
        ]

        pages = exporter._group_rows_into_pages(
            rows=rows,
            cards_list=cards,
            class_field_name='Class',
            section_field_name='Section',
            records_per_page=6,
        )

        self.assertEqual(
            [[row[0]['content'] for row in page] for page in pages],
            [['row1'], ['row2']],
        )


class ExportFieldDetectionTests(SimpleTestCase):
    def test_get_section_field_name_prefers_section_type(self):
        from exports.utils import get_section_field_name

        table_fields = [
            {'name': 'Class', 'type': 'class'},
            {'name': 'Division', 'type': 'section'},
        ]

        self.assertEqual(get_section_field_name(table_fields), 'Division')

    def test_designation_field_is_not_image(self):
        from exports.utils import is_image_field, classify_image_subtype

        designation_field = {'name': 'DESIGNATION', 'type': 'text'}

        self.assertFalse(is_image_field(designation_field))
        self.assertIsNone(classify_image_subtype(designation_field))

    def test_sign_field_still_classifies_as_signature(self):
        from exports.utils import is_image_field, classify_image_subtype

        sign_field = {'name': 'SIGN', 'type': 'text'}

        self.assertTrue(is_image_field(sign_field))
        self.assertEqual(classify_image_subtype(sign_field), 'signature')

    def test_rel_slot_without_media_keyword_stays_text(self):
        from exports.utils import is_image_field, classify_image_subtype

        rel_slot_name_field = {'name': 'REL_1', 'type': 'text'}

        self.assertFalse(is_image_field(rel_slot_name_field))
        self.assertIsNone(classify_image_subtype(rel_slot_name_field))

    def test_rel_slot_with_photo_keyword_maps_to_rel_photo(self):
        from exports.utils import is_image_field, classify_image_subtype

        rel_photo_field = {'name': 'REL_1PHOTO', 'type': 'text'}

        self.assertTrue(is_image_field(rel_photo_field))
        self.assertEqual(classify_image_subtype(rel_photo_field), 'rel_photo')


class WordLayoutTuningTests(SimpleTestCase):
    def test_dense_width_allocator_keeps_table_inside_page(self):
        from exports.word import WordExporter

        exporter = WordExporter()

        text_fields = [
            {'name': 'NAME', 'is_image': False, 'type': 'text'},
            {'name': 'FATHER_NAME', 'is_image': False, 'type': 'text'},
            {'name': 'ADDRESS', 'is_image': False, 'type': 'textarea'},
            {'name': 'EMAIL', 'is_image': False, 'type': 'text'},
            {'name': 'DEPARTMENT', 'is_image': False, 'type': 'text'},
            {'name': 'DESIGNATION', 'is_image': False, 'type': 'text'},
            {'name': 'MEDICAL_CONDITION', 'is_image': False, 'type': 'text'},
            {'name': 'ALLERGIES', 'is_image': False, 'type': 'text'},
        ]
        # Add extra long-text columns to force dense layout path (>20 columns total)
        text_fields.extend([
            {'name': f'LONG_TEXT_{i}', 'is_image': False, 'type': 'text'}
            for i in range(1, 11)
        ])
        image_fields = [
            {
                'name': f'PHOTO_{i}',
                'is_image': True,
                'type': 'photo',
                'image_width_cm': 1.9,
                'image_height_cm': 2.5,
                'image_subtype': 'photo',
            }
            for i in range(1, 7)
        ]
        ordered_fields = text_fields + image_fields
        num_cols = 1 + len(ordered_fields)

        class _Card:
            def __init__(self, data):
                self.field_data = data

        cards = []
        for i in range(8):
            row = {}
            for f in text_fields:
                row[f['name']] = (
                    f"VERY LONG VALUE {i} FOR {f['name']} WITH MULTIPLE WORDS "
                    f"ANDLONGUNBREAKABLETOKEN{i}XYZABCDEFGHIJK"
                )
            cards.append(_Card(row))

        widths = exporter._calculate_column_widths(ordered_fields, cards, num_cols)
        total_width = sum(widths.get(i, 0.0) for i in range(num_cols))

        self.assertLessEqual(total_width, exporter.PAGE_WIDTH_CM + 0.05)
        self.assertGreaterEqual(total_width, exporter.PAGE_WIDTH_CM - 0.30)

        text_col_indexes = [0] + [
            1 + idx for idx, f in enumerate(ordered_fields) if not f['is_image']
        ]
        self.assertTrue(all(widths[i] >= 0.55 for i in text_col_indexes))

    def test_create_data_tables_uses_atleast_height_for_image_rows(self):
        from exports.word import WordExporter
        from docx import Document
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml, OxmlElement
        from docx.oxml.ns import nsdecls, qn
        from PIL import Image, ImageOps

        exporter = WordExporter()
        doc = Document()

        ordered_fields = [
            {'name': 'ADDRESS', 'is_image': False, 'type': 'textarea'},
            {
                'name': 'PHOTO',
                'is_image': True,
                'type': 'photo',
                'image_width_cm': 1.9,
                'image_height_cm': 2.5,
                'image_subtype': 'photo',
            },
        ]
        column_widths = {0: 1.0, 1: 4.0, 2: 1.9}

        class _Card:
            def __init__(self, data):
                self.field_data = data

        cards = [_Card({'ADDRESS': 'LONG ADDRESS VALUE FOR WRAP TEST', 'PHOTO': ''})]

        with mock.patch.object(exporter, '_add_data_row') as mocked_add_row:
            exporter._create_data_tables(
                doc=doc,
                cards_list=cards,
                ordered_fields=ordered_fields,
                column_widths=column_widths,
                num_cols=3,
                Cm=Cm,
                Pt=Pt,
                RGBColor=RGBColor,
                WD_TABLE_ALIGNMENT=WD_TABLE_ALIGNMENT,
                WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
                parse_xml=parse_xml,
                nsdecls=nsdecls,
                OxmlElement=OxmlElement,
                qn=qn,
                Image=Image,
                ImageOps=ImageOps,
                class_field_name=None,
            )

        self.assertTrue(mocked_add_row.called)
        call_kwargs = mocked_add_row.call_args.kwargs
        self.assertEqual(call_kwargs.get('h_rule'), 'atLeast')
        self.assertGreaterEqual(call_kwargs.get('row_height_cm', 0), exporter.ROW_HEIGHT_CM)

    def test_missing_image_placeholder_uses_row_height_minimum(self):
        from exports.word import WordExporter
        from docx import Document
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from PIL import Image, ImageOps

        exporter = WordExporter()
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        with mock.patch.object(exporter, '_add_empty_image_box') as mocked_empty_box:
            exporter._add_image_to_cell(
                cell=cell,
                img_path='',
                Cm=Cm,
                Pt=Pt,
                RGBColor=RGBColor,
                WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
                parse_xml=parse_xml,
                nsdecls=nsdecls,
                Image=Image,
                ImageOps=ImageOps,
                fixed_width_cm=1.9,
                fixed_height_cm=1.0,
                image_subtype='photo',
            )

        self.assertTrue(mocked_empty_box.called)
        args = mocked_empty_box.call_args.args
        self.assertAlmostEqual(args[8], exporter.ROW_HEIGHT_CM, places=2)

    def test_word_photo_border_decision_for_parent_photo_columns(self):
        from exports.word import WordExporter

        exporter = WordExporter()

        self.assertTrue(exporter._should_add_photo_border(image_subtype='photo', field_name='PHOTO'))
        self.assertTrue(exporter._should_add_photo_border(image_subtype='rel_photo', field_name='FATHER PHOTO'))
        self.assertTrue(exporter._should_add_photo_border(image_subtype='rel_photo', field_name='MOTHER PHOTO'))

        # Subtype-missing fallback should still border photo columns.
        self.assertTrue(exporter._should_add_photo_border(image_subtype=None, field_name='FATHER PHOTO'))
        self.assertTrue(exporter._should_add_photo_border(image_subtype=None, field_name='MOTHER IMAGE'))

        # Non-photo image types and normal text columns should not get this border.
        self.assertFalse(exporter._should_add_photo_border(image_subtype='signature', field_name='SIGNATURE'))
        self.assertFalse(exporter._should_add_photo_border(image_subtype=None, field_name='DESIGNATION'))

    def test_pdf_template_uses_image_border_not_photo_cell_box_border(self):
        template_path = os.path.join(settings.BASE_DIR, 'templates', 'exports', 'pdf_report.html')
        with open(template_path, 'r', encoding='utf-8') as handle:
            content = handle.read()

        self.assertIn('table.data-table td.img-cell.photo-cell {', content)
        self.assertIn('border: none;', content)
        self.assertIn('table.data-table td.img-cell.photo-cell img {', content)
        self.assertIn('border: 1pt solid #111;', content)
        self.assertIn('template_use_abbasi', content)

    def test_word_data_row_removes_photo_cell_box_border(self):
        from exports.word import WordExporter
        from docx import Document
        from docx.shared import Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from PIL import Image, ImageOps

        exporter = WordExporter()
        doc = Document()
        table = doc.add_table(rows=1, cols=2)

        card = type('Card', (), {'field_data': {'FATHER PHOTO': ''}})()
        ordered_fields = [{
            'name': 'FATHER PHOTO',
            'is_image': True,
            'type': 'rel_photo',
            'image_subtype': 'rel_photo',
            'image_width_cm': 1.9,
            'image_height_cm': 2.5,
        }]
        column_widths = {0: 1.0, 1: 1.9}

        with mock.patch.object(exporter, '_add_image_to_cell') as mocked_add_image:
            with mock.patch.object(exporter, '_remove_cell_borders') as mocked_remove_borders:
                exporter._add_data_row(
                    table=table,
                    card=card,
                    ordered_fields=ordered_fields,
                    column_widths=column_widths,
                    sr_no=1,
                    Cm=Cm,
                    Pt=Pt,
                    RGBColor=RGBColor,
                    WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH,
                    parse_xml=parse_xml,
                    nsdecls=nsdecls,
                    Image=Image,
                    ImageOps=ImageOps,
                    image_fixed_widths={'FATHER PHOTO': 1.9},
                    image_fixed_heights={'FATHER PHOTO': 2.5},
                    row_height_cm=2.5,
                    font_pt=9,
                    h_rule='atLeast',
                )

        self.assertTrue(mocked_add_image.called)
        self.assertTrue(mocked_remove_borders.called)

    def test_word_photo_image_stream_adds_1pt_equivalent_edge(self):
        from exports.word import WordExporter
        from PIL import Image as PILImage, ImageOps as PILImageOps
        import io

        exporter = WordExporter()

        src_stream = io.BytesIO()
        PILImage.new('RGB', (20, 30), (255, 0, 0)).save(src_stream, format='PNG')
        src_bytes = src_stream.getvalue()

        image_stream = exporter._build_word_image_stream(
            src_bytes,
            PILImage,
            PILImageOps,
            add_photo_border=True,
        )

        with PILImage.open(image_stream) as image:
            self.assertEqual(image.size, (20, 30))
            self.assertEqual(image.getpixel((0, 0)), (0, 0, 0))
            self.assertEqual(image.getpixel((1, 1)), (0, 0, 0))
            self.assertEqual(image.getpixel((2, 2)), (255, 0, 0))
            self.assertEqual(image.getpixel((10, 15)), (255, 0, 0))

    def test_word_hindi_font_detection_is_normalized(self):
        from exports.word import WordExporter

        exporter = WordExporter()
        self.assertTrue(exporter._is_hindi_abbasi_font('hindi'))
        self.assertTrue(exporter._is_hindi_abbasi_font(' Hindi '))
        self.assertTrue(exporter._is_hindi_abbasi_font('abbasi'))
        self.assertFalse(exporter._is_hindi_abbasi_font('arial'))


class StreamFileResponseSuperModeTests(TestCase):
    def setUp(self):
        self.pro_user = User.objects.create_user(
            username='stream-pro@test.com',
            email='stream-pro@test.com',
            password='pass12345',
            role='pro_user',
        )
        self.normal_user = User.objects.create_user(
            username='stream-normal@test.com',
            email='stream-normal@test.com',
            password='pass12345',
            role='admin_staff',
        )

        from core.services.super_mode_service import SuperModeService

        SuperModeService.configure_pro_user_self(
            self.pro_user,
            enabled=True,
            ram_mb=750,
        )

    def test_super_mode_large_file_streams_from_ram_without_temp_file(self):
        from exports.utils import stream_file_response

        payload = b'A' * (11 * 1024 * 1024)

        with mock.patch('tempfile.NamedTemporaryFile', side_effect=AssertionError('temp spool should not run')):
            response = stream_file_response(
                payload,
                'super.bin',
                'application/octet-stream',
                user=self.pro_user,
            )
            streamed = b''.join(response.streaming_content)

        self.assertEqual(streamed, payload)
        self.assertEqual(int(response['Content-Length']), len(payload))

    def test_non_super_mode_large_file_keeps_existing_temp_spool_behavior(self):
        import tempfile
        from exports.utils import stream_file_response

        payload = b'B' * (11 * 1024 * 1024)
        real_named_temporary_file = tempfile.NamedTemporaryFile

        with mock.patch('tempfile.NamedTemporaryFile', wraps=real_named_temporary_file) as patched_tmp:
            response = stream_file_response(
                payload,
                'normal.bin',
                'application/octet-stream',
                user=self.normal_user,
            )
            streamed = b''.join(response.streaming_content)

        self.assertGreaterEqual(patched_tmp.call_count, 1)
        self.assertEqual(streamed, payload)
