"""
Export settings & template API views.
Split from base.py for maintainability.
"""
import json
import logging
import base64
import html as py_html
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
from django.contrib.auth.decorators import login_required
from django.utils.html import strip_tags

from ..models import SystemSettings
from ..services.activity_service import ActivityService
from ..services.permission_service import PermissionService, require_any_admin

logger = logging.getLogger(__name__)

_MAX_TEMPLATE_HTML_LEN = 250000
_MAX_TEMPLATE_IMPORT_DOC_SIZE = 12 * 1024 * 1024


def _template_has_content(instructions_html: str) -> bool:
    raw = str(instructions_html or '')
    plain = py_html.unescape(strip_tags(raw or '')).strip()
    has_image = '<img' in raw.lower()
    return bool(plain or has_image)


def _docx_escape_text(value: str) -> str:
    return py_html.escape(str(value or ''), quote=False).replace('\n', '<br>').replace('\t', '&emsp;')


def _docx_run_to_html(run, related_parts) -> str:
    from docx.oxml.ns import qn

    chunks = []
    run_text = str(getattr(run, 'text', '') or '')
    if run_text:
        text_html = _docx_escape_text(run_text)
        if run.bold:
            text_html = f'<strong>{text_html}</strong>'
        if run.italic:
            text_html = f'<em>{text_html}</em>'
        if run.underline:
            text_html = f'<u>{text_html}</u>'
        chunks.append(text_html)

    for blip in run._r.xpath('.//a:blip'):
        rel_id = blip.get(qn('r:embed'))
        if not rel_id:
            continue
        image_part = related_parts.get(rel_id)
        if not image_part:
            continue
        content_type = str(getattr(image_part, 'content_type', '') or '').lower()
        if not content_type.startswith('image/'):
            continue
        blob = getattr(image_part, 'blob', b'') or b''
        if not blob:
            continue
        data_b64 = base64.b64encode(blob).decode('ascii')
        chunks.append(f'<img src="data:{content_type};base64,{data_b64}" alt="Imported image">')

    return ''.join(chunks)


def _docx_paragraph_to_html(paragraph, related_parts) -> str:
    chunks = []
    for run in paragraph.runs:
        run_html = _docx_run_to_html(run, related_parts)
        if run_html:
            chunks.append(run_html)

    body_html = ''.join(chunks).strip() or '<br>'
    align_map = {
        0: 'left',
        1: 'center',
        2: 'right',
        3: 'justify',
    }
    align = align_map.get(getattr(paragraph, 'alignment', None))
    style_attr = f' style="text-align:{align};"' if align else ''
    return f'<p{style_attr}>{body_html}</p>'


def _docx_table_to_html(table, related_parts) -> str:
    row_html = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_blocks = []
            for p in cell.paragraphs:
                cell_blocks.append(_docx_paragraph_to_html(p, related_parts))
            cells.append(f'<td>{"".join(cell_blocks) or "<p><br></p>"}</td>')
        row_html.append(f'<tr>{"".join(cells)}</tr>')
    return f'<table><tbody>{"".join(row_html)}</tbody></table>'


def _docx_to_editor_html(uploaded_file) -> str:
    from docx import Document

    document = Document(uploaded_file)
    related_parts = document.part.related_parts

    paragraph_map = {p._p: p for p in document.paragraphs}
    table_map = {t._tbl: t for t in document.tables}

    blocks = []
    for child in document.element.body.iterchildren():
        tag = str(getattr(child, 'tag', '') or '')
        if tag.endswith('}p'):
            paragraph = paragraph_map.get(child)
            if paragraph is not None:
                blocks.append(_docx_paragraph_to_html(paragraph, related_parts))
            continue

        if tag.endswith('}tbl'):
            table = table_map.get(child)
            if table is not None:
                blocks.append(_docx_table_to_html(table, related_parts))

    editor_html = ''.join(blocks).strip() or '<p><br></p>'
    if len(editor_html) > _MAX_TEMPLATE_HTML_LEN:
        raise ValueError('Imported document is too large for the editor. Please reduce content and try again.')
    return editor_html


# =========================================================================
# EXPORT SETTINGS API
# =========================================================================

@login_required
@require_any_admin
@require_http_methods(['GET'])
def api_export_settings_get(request):
    """GET /api/export-settings/ — fetch export footer messages (admin only)."""
    data = SystemSettings.get_export_settings()
    return JsonResponse({'success': True, 'data': data})


@login_required
@require_any_admin
@require_http_methods(['POST'])
def api_export_settings_update(request):
    """POST /api/export-settings/update/ — update export footer messages (super admin / admin staff only)."""
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'success': False, 'message': 'Invalid JSON'}, status=400)

    updated = []
    MAX_SETTING_VALUE_LEN = 1000
    for key in SystemSettings.EXPORT_DEFAULTS:
        if key in body:
            val = body[key].strip() if isinstance(body[key], str) else str(body[key]).strip()
            if len(val) > MAX_SETTING_VALUE_LEN:
                return JsonResponse({'success': False, 'message': f'{key} exceeds maximum length of {MAX_SETTING_VALUE_LEN} characters'}, status=400)
            SystemSettings.set_value(key, val)
            updated.append(key)

    if not updated:
        return JsonResponse({'success': False, 'message': 'No valid fields provided'}, status=400)

    return JsonResponse({'success': True, 'message': 'Export settings updated successfully', 'updated': updated})


# =========================================================================
# EXPORT TEMPLATES API
# =========================================================================

@login_required
@require_http_methods(['GET'])
def api_export_templates_list(request):
    """
    GET /api/export-templates/ — list export templates for download modals.

    Access policy:
    - Any admin (super_admin/admin_staff) can always read.
    - Client/client_staff can read only if they have bulk-download permission.
    """
    if not (
        PermissionService.is_any_admin(request.user)
        or PermissionService.can_bulk_download(request.user)
    ):
        return JsonResponse(
            {'success': False, 'message': 'Permission denied: You do not have bulk download access'},
            status=403,
        )

    from core.models import ExportTemplate
    templates = ExportTemplate.get_all_as_choices()
    return JsonResponse({'success': True, 'templates': templates})


@login_required
@require_any_admin
@require_http_methods(['POST'])
def api_export_template_import_doc(request):
    """POST /api/export-templates/import-doc/ — import .docx into rich editor HTML."""
    uploaded = request.FILES.get('file')
    if not uploaded:
        return JsonResponse({'success': False, 'message': 'Please upload a .docx file'}, status=400)

    file_name = str(getattr(uploaded, 'name', '') or '').strip().lower()
    if not file_name.endswith('.docx'):
        return JsonResponse({'success': False, 'message': 'Only .docx files are supported'}, status=400)

    file_size = int(getattr(uploaded, 'size', 0) or 0)
    if file_size <= 0:
        return JsonResponse({'success': False, 'message': 'Uploaded file is empty'}, status=400)
    if file_size > _MAX_TEMPLATE_IMPORT_DOC_SIZE:
        max_mb = int(_MAX_TEMPLATE_IMPORT_DOC_SIZE / (1024 * 1024))
        return JsonResponse({'success': False, 'message': f'File too large. Maximum size is {max_mb} MB.'}, status=400)

    try:
        editor_html = _docx_to_editor_html(uploaded)
    except ValueError as exc:
        return JsonResponse({'success': False, 'message': str(exc)}, status=400)
    except Exception:
        logger.exception('DOCX import failed for export template editor')
        return JsonResponse({'success': False, 'message': 'Could not import this Word document'}, status=400)

    return JsonResponse({'success': True, 'html': editor_html})


@login_required
@require_any_admin
@require_http_methods(['POST'])
def api_export_template_create(request):
    """POST /api/export-templates/create/ — create a new export template."""
    from core.models import ExportTemplate
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'success': False, 'message': 'Invalid JSON'}, status=400)

    name = (body.get('name') or '').strip()
    instructions = str(body.get('instructions') or '').strip()
    is_default = bool(body.get('is_default', False))
    font_name = (body.get('font_name') or 'arial').strip().lower()
    is_bold = bool(body.get('is_bold', False))

    if font_name not in ('arial', 'hindi'):
        font_name = 'arial'

    if not name:
        return JsonResponse({'success': False, 'message': 'Template name is required'}, status=400)
    if not _template_has_content(instructions):
        return JsonResponse({'success': False, 'message': 'Footer content is required'}, status=400)
    if len(instructions) > _MAX_TEMPLATE_HTML_LEN:
        return JsonResponse({'success': False, 'message': f'Footer content must be {_MAX_TEMPLATE_HTML_LEN} characters or less'}, status=400)
    if len(name) > 100:
        return JsonResponse({'success': False, 'message': 'Name must be 100 characters or less'}, status=400)

    if ExportTemplate.objects.filter(name__iexact=name).exists():
        return JsonResponse({'success': False, 'message': 'A template with this name already exists'}, status=400)

    tpl = ExportTemplate.objects.create(name=name, instructions=instructions, is_default=is_default, font_name=font_name, is_bold=is_bold)
    return JsonResponse({'success': True, 'message': 'Template created', 'template': {
        'id': tpl.id, 'name': tpl.name, 'instructions': tpl.instructions,
        'font_name': tpl.font_name, 'is_bold': tpl.is_bold, 'is_default': tpl.is_default
    }})


@login_required
@require_any_admin
@require_http_methods(['POST'])
def api_export_template_update(request, template_id):
    """POST /api/export-templates/<id>/update/ — update an export template."""
    from core.models import ExportTemplate
    try:
        tpl = ExportTemplate.objects.get(id=template_id)
    except ExportTemplate.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Template not found'}, status=404)

    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'success': False, 'message': 'Invalid JSON'}, status=400)

    name = (body.get('name') or '').strip()
    instructions = str(body.get('instructions') or '').strip()
    is_default = body.get('is_default')
    font_name = body.get('font_name')
    is_bold = body.get('is_bold')

    if name:
        if len(name) > 100:
            return JsonResponse({'success': False, 'message': 'Name must be 100 characters or less'}, status=400)
        if ExportTemplate.objects.filter(name__iexact=name).exclude(pk=tpl.pk).exists():
            return JsonResponse({'success': False, 'message': 'A template with this name already exists'}, status=400)
        tpl.name = name
    if 'instructions' in body:
        if len(instructions) > _MAX_TEMPLATE_HTML_LEN:
            return JsonResponse({'success': False, 'message': f'Footer content must be {_MAX_TEMPLATE_HTML_LEN} characters or less'}, status=400)
        if not _template_has_content(instructions):
            return JsonResponse({'success': False, 'message': 'Footer content is required'}, status=400)
        tpl.instructions = instructions
    if is_default is not None:
        tpl.is_default = bool(is_default)
    if font_name is not None:
        fn = str(font_name).strip().lower()
        if fn in ('arial', 'hindi'):
            tpl.font_name = fn
    if is_bold is not None:
        tpl.is_bold = bool(is_bold)
    tpl.save()

    return JsonResponse({'success': True, 'message': 'Template updated', 'template': {
        'id': tpl.id, 'name': tpl.name, 'instructions': tpl.instructions,
        'font_name': tpl.font_name, 'is_bold': tpl.is_bold, 'is_default': tpl.is_default
    }})


@login_required
@require_any_admin
@require_http_methods(['POST'])
def api_export_template_delete(request, template_id):
    """POST /api/export-templates/<id>/delete/ — delete an export template."""
    if not PermissionService.is_super_admin(request.user):
        return JsonResponse({'success': False, 'message': 'Only super admin can delete export templates.'}, status=403)

    from core.models import ExportTemplate
    try:
        tpl = ExportTemplate.objects.get(id=template_id)
    except ExportTemplate.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Template not found'}, status=404)

    template_name = tpl.name
    template_id_value = tpl.id
    tpl.delete()

    try:
        ActivityService.log(
            'settings_update',
            f'Export template "{template_name}" deleted',
            request=request,
            target_model='ExportTemplate',
            target_id=template_id_value,
            target_name=template_name,
        )
    except Exception:
        logger.exception('Failed to log export template deletion')

    return JsonResponse({'success': True, 'message': 'Template deleted'})
