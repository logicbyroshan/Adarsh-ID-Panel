"""
Word Export — Images Mixin

Image embedding, VML conversion, and empty placeholder rendering
for Word document generation.
"""
import logging
import re
from io import BytesIO

from django.core.files.storage import default_storage

from .utils import is_valid_image_path

logger = logging.getLogger(__name__)


class WordImagesMixin:
    """Mixin providing image handling methods for Word exports."""

    # Border is intentionally limited to portrait photos only.
    BORDERED_IMAGE_SUBTYPES = {'photo', 'rel_photo', 'mother_photo', 'father_photo'}
    PHOTO_BORDER_PX = 2  # ~1pt visual thickness on typical display scaling
    PHOTO_BORDER_COLOR = (0, 0, 0)
    WORD_BORDER_PT = 1.0

    def _should_add_photo_border(self, image_subtype=None, field_name=None):
        """Return True when the image should be rendered with a 1pt border."""
        subtype = str(image_subtype or '').strip().lower()
        if subtype in self.BORDERED_IMAGE_SUBTYPES:
            return True

        # Fallback: if subtype is missing, infer from field name safely.
        name = str(field_name or '').strip().lower()
        if not name:
            return False

        if re.search(r'\b(?:father|mother)\b.*\b(?:photo|image|pic|picture)\b', name):
            return True
        if re.search(r'\b(?:photo|image|pic|picture)\b', name) and not re.search(r'\b(?:signature|sign|barcode|qr)\b', name):
            return True
        return False

    @staticmethod
    def _build_word_image_stream(img_data, Image, ImageOps, add_photo_border=False):
        """Return a BytesIO stream for Word embedding.

        Some Word compatibility renderers ignore VML stroke on image shapes.
        For photo-like columns, add a 1pt-equivalent inner border in the
        bitmap (without changing dimensions) so border remains visible.
        """
        if not add_photo_border:
            return BytesIO(img_data)

        with Image.open(BytesIO(img_data)) as src_img:
            src_img.load()
            if src_img.mode != 'RGB':
                src_img = src_img.convert('RGB')
            bordered = src_img.copy()
            w, h = bordered.size
            if w >= 2 and h >= 2:
                px = bordered.load()
                border_px = WordImagesMixin.PHOTO_BORDER_PX
                edge = WordImagesMixin.PHOTO_BORDER_COLOR
                max_t = min(border_px, w // 2, h // 2)
                for t in range(max_t):
                    for x in range(t, w - t):
                        px[x, t] = edge
                        px[x, h - 1 - t] = edge
                    for y in range(t, h - t):
                        px[t, y] = edge
                        px[w - 1 - t, y] = edge
            out_stream = BytesIO()
            bordered.save(out_stream, format='PNG', optimize=True)
            out_stream.seek(0)
            return out_stream

    def _add_image_to_cell(self, cell, img_path, Cm, Pt, RGBColor,
                           WD_ALIGN_PARAGRAPH, parse_xml, nsdecls, Image, ImageOps,
                           fixed_width_cm=None, fixed_height_cm=None,
                           image_subtype=None, field_name=None):
        """Add an image to a cell using VML for Word 97-2003 compatibility.
        
        Uses VML (Vector Markup Language) instead of DrawingML to ensure
        images are visible in Normal/Draft view and compatible with
        Word 97-2003 format.
        
        For missing/pending images: draws an empty bordered rectangle
        (no placeholder image, no text).
        """
        if fixed_width_cm is None:
            fixed_width_cm = self.IMAGE_DEFAULT_WIDTH_CM
        if fixed_height_cm is None:
            fixed_height_cm = self.IMAGE_HEIGHT_CM

        self._set_cell_margins(cell, parse_xml, nsdecls, 0, 0, 0, 0)
        self._set_cell_vertical_align(cell, parse_xml, nsdecls)
        
        if is_valid_image_path(img_path):
            try:
                if default_storage.exists(img_path):
                    with default_storage.open(img_path, 'rb') as img_file:
                        img_data = img_file.read()
                        
                        if img_data and len(img_data) >= 100:
                            # Validate image and keep original bytes (no recompression).
                            with Image.open(BytesIO(img_data)) as verify_img:
                                verify_img.verify()

                            add_photo_border = self._should_add_photo_border(
                                image_subtype=image_subtype,
                                field_name=field_name,
                            )
                            img_stream = self._build_word_image_stream(
                                img_data,
                                Image,
                                ImageOps,
                                add_photo_border=add_photo_border,
                            )
                            
                            para = cell.paragraphs[0]
                            run = para.add_run()
                            # Keep layout fixed while preserving source quality
                            # (original bytes, no JPEG recompression step).
                            inline_shape = run.add_picture(
                                img_stream,
                                height=Cm(fixed_height_cm),
                                width=Cm(fixed_width_cm)
                            )
                            img_stream.close()
                            self._convert_to_vml(
                                run,
                                inline_shape,
                                add_border=False,
                            )
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._set_para_spacing(para, parse_xml, nsdecls)
                            return
                        else:
                            logger.warning("Word export: Image too small (%d bytes): %s", len(img_data) if img_data else 0, img_path)
                else:
                    logger.warning("Word export: Image file not found by storage: %s", img_path)
            except Exception as e:
                logger.warning("Word export: Image load error for %s: %s", img_path, e)
        
        # Missing/pending image → draw an empty bordered rectangle
        # (no placeholder image, no text — just empty space with a border).
        # Keep placeholder at least ROW_HEIGHT_CM so image rows remain visually
        # consistent even when the source image is missing.
        placeholder_h = max(fixed_height_cm or 0, getattr(self, 'ROW_HEIGHT_CM', 2.5))
        self._add_empty_image_box(cell, Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                                   parse_xml, nsdecls, fixed_width_cm, placeholder_h)

    def _add_empty_image_box(self, cell, Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                              parse_xml, nsdecls, fixed_width_cm, fixed_height_cm=None):
        """Draw an empty bordered rectangle for missing/pending images.
        
        Uses a 1-row, 1-col inner table with fixed dimensions and
        a thin black border to represent an empty image placeholder.
        """
        from docx.enum.table import WD_TABLE_ALIGNMENT

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_para_spacing(para, parse_xml, nsdecls)

        # Clear any default text
        for run in para.runs:
            run.clear()

        # Create a tiny inline table as the bordered box
        # We use VML rectangle via a Run for better compatibility
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement

        # Width/Height in EMU (1 cm = 360000 EMU)
        effective_h = fixed_height_cm if fixed_height_cm else self.IMAGE_HEIGHT_CM
        box_w_emu = int(fixed_width_cm * 360000)
        box_h_emu = int(effective_h * 360000)
        box_w_pt = fixed_width_cm * 28.3465
        box_h_pt = effective_h * 28.3465

        # Create VML rectangle shape directly
        run = para.add_run()
        pict_xml = (
            '<w:pict '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office">'
            f'<v:rect style="width:{box_w_pt:.1f}pt;height:{box_h_pt:.1f}pt" '
            f'filled="f" strokecolor="#000000" strokeweight="{self.WORD_BORDER_PT:.1f}pt">'
            '</v:rect></w:pict>'
        )
        from lxml import etree
        pict_elem = etree.fromstring(pict_xml)
        run._r.append(pict_elem)
    
    def _convert_to_vml(self, run, inline_shape, add_border=False):
        """Convert an inline DrawingML image to VML for backward compatibility.
        
        VML images are visible in Word's Normal/Draft view and compatible
        with Word 97-2003 format (.doc).  This replaces the <w:drawing>
        element with a <w:pict> VML element referencing the same image
        relationship.
        """
        from lxml import etree  # type: ignore[attr-defined]
        from docx.oxml.ns import qn
        
        # Get dimensions from inline shape (EMU → points)
        width_pt = inline_shape.width / 914400.0 * 72.0
        height_pt = inline_shape.height / 914400.0 * 72.0
        
        # Extract relationship ID from DrawingML blip element
        drawing_elem = run._r.find(qn('w:drawing'))
        blip = drawing_elem.find('.//' + qn('a:blip'))
        rId = blip.get(qn('r:embed'))
        
        # Remove the DrawingML element from the run
        run._r.remove(drawing_elem)
        
        # Create VML <w:pict> element (universally compatible).
        # Use <v:rect> so stroke is consistently rendered around the image.
        border_attrs = (
            f'stroked="t" strokecolor="#000000" strokeweight="{self.WORD_BORDER_PT:.1f}pt"'
            if add_border else
            'stroked="f"'
        )
        vml_xml = (
            '<w:pict '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<v:rect {border_attrs} '
            f'style="width:{width_pt:.1f}pt;height:{height_pt:.1f}pt">'
            f'<v:imagedata r:id="{rId}" o:title=""/>'
            '</v:rect></w:pict>'
        )
        pict_elem = etree.fromstring(vml_xml)
        run._r.append(pict_elem)
