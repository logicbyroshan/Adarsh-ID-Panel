"""
Export Service — Phase 11

Handles:
  - PDF export (chunked, memory-safe, WeasyPrint)
  - DOCX export (python-docx)
  - XLSX export
  - ZIP image export with rename patterns

All exports:
  - Apply scope (selected / filtered / all)
  - Sort by Class → Section → Name
  - Transition APPROVED → DOWNLOADED on success
  - Generate WorkflowHistory + AuditLog entries
  - Are driven by ExportSession + ExportResult + ExportArtifact records
"""
import os
import io
import uuid
import zipfile
import logging
import tempfile
import shutil
from io import BytesIO
from typing import List, Optional

from django.conf import settings
from django.utils import timezone
from django.db import transaction

from apps.cards.models import Card
from apps.fields.models import Field, FieldType
from apps.mediafiles.models import MediaFile, MediaVariant
from apps.mediafiles.storage.factory import StorageFactory
from apps.workflow.constants import WorkflowState, WorkflowAction
from apps.workflow.services import WorkflowService
from apps.workflow.models import WorkflowHistory
from apps.auditlogs.models import AuditLog
from apps.exports.constants import ExportStatus, ExportType, PageBreak, XlsxFieldScope, ExportScope
from apps.exports.models import ExportSession, ExportResult, ExportArtifact
from apps.exports.placeholder import PlaceholderParser
from apps.exports.sorting import sort_cards
from apps.exports.rename import RenameEngine

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────

CHUNK_SIZE = 500  # records per rendering chunk


def _get_image_bytes(card: Card, field: Field, use_thumbnail: bool = True) -> Optional[bytes]:
    """Retrieve image bytes from storage for a given card/field combo."""
    try:
        from apps.mediafiles.models import MediaReference
        ref = MediaReference.objects.select_related('media_file').filter(
            card=card, field=field
        ).first()
        if not ref:
            return None
        media_file = ref.media_file
        if media_file.is_deleted:
            return None

        stored_path = media_file.stored_name
        if use_thumbnail:
            thumb = media_file.variants.filter(variant_name='thumbnail').first()
            if thumb:
                stored_path = thumb.stored_name

        storage = StorageFactory.get_storage()
        return storage.read(stored_path)
    except Exception as e:
        logger.warning(f"Failed to retrieve image for card {card.id} field {field.id}: {e}")
        return None


def _resolve_cards(table, options: dict) -> list:
    """
    Resolve the card queryset based on export scope rules:
      - If card_ids provided → selected cards only (excluding DELETED)
      - Else if status filter → filtered by that exact status
      - Else → all non-DELETED cards in table
    """
    from apps.workflow.constants import WorkflowState
    card_ids = options.get('card_ids', [])
    status_filter = options.get('status')

    qs = Card.objects.filter(table=table)

    if card_ids:
        # Respect explicit selection but never include DELETED
        qs = qs.filter(id__in=card_ids).exclude(status=WorkflowState.DELETED)
    elif status_filter:
        qs = qs.filter(status=status_filter)
    else:
        # Exclude soft-deleted cards
        qs = qs.exclude(status=WorkflowState.DELETED)

    return list(qs.select_related('table', 'organization'))


def _mark_downloaded(cards: list, user, export_session: ExportSession):
    """Transition APPROVED cards to DOWNLOADED. Called after any successful export."""
    from apps.workflow.constants import WorkflowState
    for card in cards:
        # Re-fetch status to avoid stale in-memory value
        card.refresh_from_db(fields=['status'])
        if card.status == WorkflowState.APPROVED:
            try:
                WorkflowService.transition_card(
                    card=card,
                    action=WorkflowAction.DOWNLOAD,
                    user=user,
                    reason=f"Exported via ExportSession {export_session.id}"
                )
            except Exception as e:
                logger.warning(f"Could not mark card {card.id} as DOWNLOADED: {e}")


def _build_card_context(card: Card, fields) -> dict:
    """Build the template context for a single card."""
    ctx = PlaceholderParser.build_context(card, fields)
    ctx['display_id'] = card.display_id or ''
    return ctx


def _store_artifact(export_session: ExportSession, data: bytes, filename: str, mime_type: str) -> ExportArtifact:
    """Store the generated file in media storage and create an ExportArtifact record."""
    storage = StorageFactory.get_storage()
    uid = str(uuid.uuid4())
    stored_path = f"exports/{uid[:2]}/{uid[2:4]}/{uid}_{filename}"
    storage.save(stored_path, BytesIO(data))

    return ExportArtifact.objects.create(
        export_session=export_session,
        file_name=filename,
        stored_path=stored_path,
        mime_type=mime_type,
        file_size=len(data),
    )


# ──────────────────────────────────────────────────
# PDF Renderer
# ──────────────────────────────────────────────────

class PdfRenderer:
    """
    Renders cards to PDF using WeasyPrint in 500-record chunks,
    then merges with pypdf.
    """

    @staticmethod
    def render(
        export_session: ExportSession,
        cards: list,
        fields,
        template,
        page_break: str = PageBreak.NONE,
        progress_callback=None,
    ) -> bytes:
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            raise RuntimeError("WeasyPrint is not installed. Run: pip install weasyprint")

        try:
            from pypdf import PdfWriter, PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfWriter, PdfReader
            except ImportError:
                raise RuntimeError("pypdf is not installed. Run: pip install pypdf")

        body_template = (template.body or '') if template else '<p>{{display_id}}</p>'
        field_list = list(fields)
        image_fields = [f for f in field_list if f.type == FieldType.IMAGE]

        # Sort cards
        sorted_cards = sort_cards(cards, field_list)
        total = len(sorted_cards)

        writer = PdfWriter()
        chunk_no = 0

        # Group cards for page-break logic
        def grouped_chunks():
            if page_break == PageBreak.NONE:
                for i in range(0, total, CHUNK_SIZE):
                    yield sorted_cards[i:i + CHUNK_SIZE]
            else:
                # sort further into page groups, each group becomes a chunk
                groups = {}
                class_fid = None
                section_fid = None
                for f in field_list:
                    nm = f.name.lower()
                    if 'class' in nm or 'grade' in nm:
                        class_fid = str(f.id)
                    if 'section' in nm or 'division' in nm:
                        section_fid = str(f.id)

                for card in sorted_cards:
                    if page_break == PageBreak.BY_CLASS:
                        key = str(card.data.get(class_fid, '') or '') if class_fid else 'all'
                    else:
                        key_c = str(card.data.get(class_fid, '') or '') if class_fid else ''
                        key_s = str(card.data.get(section_fid, '') or '') if section_fid else ''
                        key = f"{key_c}_{key_s}"
                    groups.setdefault(key, []).append(card)

                for key, group in groups.items():
                    for i in range(0, len(group), CHUNK_SIZE):
                        yield group[i:i + CHUNK_SIZE]

        temp_dir = tempfile.mkdtemp()
        try:
            chunk_pdfs = []
            for chunk in grouped_chunks():
                chunk_no += 1
                html_parts = []

                for card in chunk:
                    ctx = _build_card_context(card, field_list)

                    # Embed images as base64
                    for img_field in image_fields:
                        img_bytes = _get_image_bytes(card, img_field, use_thumbnail=True)
                        if img_bytes:
                            import base64
                            b64 = base64.b64encode(img_bytes).decode('utf-8')
                            ext = img_field.name.lower()
                            mime = 'image/jpeg'
                            ctx[PlaceholderParser._field_key(img_field.name)] = (
                                f'<img src="data:{mime};base64,{b64}" style="max-width:150px;max-height:150px;" />'
                            )

                    rendered = PlaceholderParser.render(body_template, ctx)

                    if page_break != PageBreak.NONE:
                        rendered += '<div style="page-break-after: always;"></div>'
                    html_parts.append(rendered)

                full_html = '<html><body>' + ''.join(html_parts) + '</body></html>'

                chunk_path = os.path.join(temp_dir, f"chunk_{chunk_no}.pdf")
                HTML(string=full_html).write_pdf(chunk_path)
                chunk_pdfs.append(chunk_path)

                processed = min(chunk_no * CHUNK_SIZE, total)
                if progress_callback:
                    pct = int(10 + (processed / max(total, 1)) * 80)
                    progress_callback(pct, f"Rendered {processed}/{total} records")

            # Merge all chunk PDFs
            for path in chunk_pdfs:
                reader = PdfReader(path)
                for page in reader.pages:
                    writer.add_page(page)

            buf = BytesIO()
            writer.write(buf)
            buf.seek(0)
            return buf.read()

        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)


# ──────────────────────────────────────────────────
# DOCX Renderer
# ──────────────────────────────────────────────────

class DocxRenderer:
    """Renders cards to DOCX using python-docx."""

    @staticmethod
    def render(
        export_session: ExportSession,
        cards: list,
        fields,
        template,
        page_break: str = PageBreak.NONE,
        progress_callback=None,
    ) -> bytes:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            raise RuntimeError("python-docx is not installed.")

        field_list = list(fields)
        image_fields = [f for f in field_list if f.type == FieldType.IMAGE]
        sorted_cards = sort_cards(cards, field_list)
        total = len(sorted_cards)

        body_template = (template.body or '{{display_id}}') if template else '{{display_id}}'

        doc = Document()
        temp_dir = tempfile.mkdtemp()
        try:
            for idx, card in enumerate(sorted_cards):
                ctx = _build_card_context(card, field_list)

                # Render text
                rendered_text = PlaceholderParser.render(body_template, ctx)
                doc.add_paragraph(rendered_text)

                # Add images
                for img_field in image_fields:
                    img_bytes = _get_image_bytes(card, img_field, use_thumbnail=False)
                    if img_bytes:
                        try:
                            # Detect extension from media file
                            from apps.mediafiles.models import MediaReference
                            ref = MediaReference.objects.filter(card=card, field=img_field).first()
                            ext = '.jpg'
                            if ref and ref.media_file.extension:
                                ext = '.' + ref.media_file.extension.lstrip('.')
                            img_path = os.path.join(temp_dir, f"{card.id}_{img_field.id}{ext}")
                            with open(img_path, 'wb') as f:
                                f.write(img_bytes)
                            doc.add_picture(img_path, width=Inches(2))
                        except Exception as e:
                            logger.warning(f"Could not add image to DOCX: {e}")

                if page_break != PageBreak.NONE and idx < total - 1:
                    doc.add_page_break()

                if progress_callback and idx % 100 == 0:
                    pct = int(10 + (idx / max(total, 1)) * 80)
                    progress_callback(pct, f"Rendered {idx + 1}/{total} records")

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)


# ──────────────────────────────────────────────────
# XLSX Renderer
# ──────────────────────────────────────────────────

class XlsxRenderer:
    @staticmethod
    def render(
        export_session: ExportSession,
        cards: list,
        fields,
        field_scope: str = XlsxFieldScope.ALL,
        progress_callback=None,
    ) -> bytes:
        import openpyxl
        field_list = [f for f in fields if not f.is_deleted]
        if field_scope == XlsxFieldScope.VISIBLE:
            # For now, treat all non-IMAGE fields as "visible"
            field_list = [f for f in field_list if f.type != FieldType.IMAGE]

        sorted_cards = sort_cards(cards, field_list)
        total = len(sorted_cards)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Export"

        # Header
        headers = ['Display ID', 'Status'] + [f.name for f in field_list]
        ws.append(headers)

        for idx, card in enumerate(sorted_cards):
            row = [card.display_id, card.status]
            for f in field_list:
                val = card.data.get(str(f.id), '')
                if f.type == FieldType.IMAGE:
                    val = '[IMAGE]' if val else ''
                row.append(str(val) if val is not None else '')
            ws.append(row)

            if progress_callback and idx % 500 == 0:
                pct = int(10 + (idx / max(total, 1)) * 80)
                progress_callback(pct, f"Written {idx + 1}/{total} rows")

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.read()


# ──────────────────────────────────────────────────
# ZIP Image Renderer
# ──────────────────────────────────────────────────

class ZipRenderer:
    @staticmethod
    def render(
        export_session: ExportSession,
        cards: list,
        fields,
        rename_pattern: str = '{display_id}',
        progress_callback=None,
    ) -> bytes:
        field_list = list(fields)
        image_fields = [f for f in field_list if f.type == FieldType.IMAGE]
        sorted_cards = sort_cards(cards, field_list)
        total = len(sorted_cards)

        buf = BytesIO()
        seen_names = {}

        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for idx, card in enumerate(sorted_cards):
                ctx = _build_card_context(card, field_list)

                for img_field in image_fields:
                    img_bytes = _get_image_bytes(card, img_field, use_thumbnail=False)
                    if not img_bytes:
                        continue

                    # Detect actual extension from media file
                    from apps.mediafiles.models import MediaReference
                    ref = MediaReference.objects.filter(card=card, field=img_field).first()
                    actual_ext = '.jpg'
                    if ref and ref.media_file.extension:
                        actual_ext = '.' + ref.media_file.extension.lstrip('.')

                    base_name = RenameEngine.build_filename(
                        rename_pattern, ctx, extension=actual_ext, index=idx
                    )
                    # Ensure uniqueness within archive
                    counter = seen_names.get(base_name, 0)
                    if counter:
                        name_part, ext = os.path.splitext(base_name)
                        base_name = f"{name_part}_{counter}{ext}"
                    seen_names[base_name] = counter + 1

                    zf.writestr(base_name, img_bytes)

                if progress_callback and idx % 100 == 0:
                    pct = int(10 + (idx / max(total, 1)) * 80)
                    progress_callback(pct, f"Zipped {idx + 1}/{total} cards")

        buf.seek(0)
        return buf.read()


# ──────────────────────────────────────────────────
# Main Export Service
# ──────────────────────────────────────────────────

class ExportService:

    @staticmethod
    @transaction.atomic
    def start_export(
        user,
        table,
        export_type: str,
        template=None,
        options: dict = None,
    ) -> ExportSession:
        """
        Create and persist an ExportSession.
        The actual rendering runs inside the Celery task (run_export_task).
        """
        options = options or {}
        org = table.organization

        session = ExportSession.objects.create(
            user=user,
            organization=org,
            table=table,
            template=template,
            export_type=export_type,
            options=options,
            status=ExportStatus.PENDING,
        )
        return session

    @staticmethod
    def process_export(
        export_session: ExportSession,
        job=None,
        progress_callback=None,
    ):
        """
        Execute the full export pipeline.
        Called from the Celery task.
        """
        export_session.status = ExportStatus.PROCESSING
        export_session.started_at = timezone.now()
        export_session.save()

        AuditLog.objects.create(
            event_type='EXPORT_START',
            actor=export_session.user,
            target_organization=export_session.organization,
            details={
                'export_session_id': str(export_session.id),
                'export_type': export_session.export_type,
                'table_id': str(export_session.table_id),
            }
        )

        try:
            table = export_session.table
            fields = Field.objects.filter(table=table, is_deleted=False).order_by('display_order', 'created_at')
            options = export_session.options or {}
            template = export_session.template

            # Resolve scope
            cards = _resolve_cards(table, options)
            export_session.record_count = len(cards)
            export_session.save(update_fields=['record_count'])

            export_type = export_session.export_type

            def _progress(pct, step):
                if progress_callback:
                    progress_callback(pct, step)

            # ── Render ──────────────────────────────────────
            if export_type == ExportType.PDF:
                page_break = options.get('page_break', PageBreak.NONE)
                data = PdfRenderer.render(
                    export_session=export_session,
                    cards=cards,
                    fields=fields,
                    template=template,
                    page_break=page_break,
                    progress_callback=_progress,
                )
                filename = f"export_{export_session.id}.pdf"
                mime = 'application/pdf'

            elif export_type == ExportType.DOCX:
                page_break = options.get('page_break', PageBreak.NONE)
                data = DocxRenderer.render(
                    export_session=export_session,
                    cards=cards,
                    fields=fields,
                    template=template,
                    page_break=page_break,
                    progress_callback=_progress,
                )
                filename = f"export_{export_session.id}.docx"
                mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            elif export_type == ExportType.XLSX:
                field_scope = options.get('field_scope', XlsxFieldScope.ALL)
                data = XlsxRenderer.render(
                    export_session=export_session,
                    cards=cards,
                    fields=fields,
                    field_scope=field_scope,
                    progress_callback=_progress,
                )
                filename = f"export_{export_session.id}.xlsx"
                mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

            elif export_type == ExportType.ZIP:
                rename_pattern = options.get('rename_pattern', '{display_id}')
                data = ZipRenderer.render(
                    export_session=export_session,
                    cards=cards,
                    fields=fields,
                    rename_pattern=rename_pattern,
                    progress_callback=_progress,
                )
                filename = f"export_{export_session.id}.zip"
                mime = 'application/zip'

            else:
                raise ValueError(f"Unknown export type: {export_type}")

            # ── Store artifact ───────────────────────────────
            artifact = _store_artifact(export_session, data, filename, mime)

            # ── Workflow: APPROVED → DOWNLOADED (all export types) ──
            user = export_session.user
            _mark_downloaded(cards, user, export_session)

            # ── Finalize session ─────────────────────────────
            completed = timezone.now()
            export_session.status = ExportStatus.COMPLETED
            export_session.completed_at = completed
            export_session.file_size = artifact.file_size
            if export_session.started_at:
                export_session.duration = (completed - export_session.started_at).total_seconds()
            export_session.save()

            AuditLog.objects.create(
                event_type='EXPORT_COMPLETE',
                actor=export_session.user,
                target_organization=export_session.organization,
                details={
                    'export_session_id': str(export_session.id),
                    'export_type': export_type,
                    'record_count': len(cards),
                    'file_size': artifact.file_size,
                }
            )

        except Exception as exc:
            export_session.status = ExportStatus.FAILED
            export_session.error_message = str(exc)
            export_session.completed_at = timezone.now()
            if export_session.started_at:
                export_session.duration = (export_session.completed_at - export_session.started_at).total_seconds()
            export_session.save()

            AuditLog.objects.create(
                event_type='EXPORT_FAILED',
                actor=export_session.user,
                target_organization=export_session.organization,
                details={
                    'export_session_id': str(export_session.id),
                    'error': str(exc),
                }
            )
            raise
