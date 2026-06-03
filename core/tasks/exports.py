import io
import uuid
import zipfile
from typing import List, Dict
from celery import shared_task
from django.conf import settings
from core.models import Job, JobLog, CardRecord, DynamicTable
from core.services.storage import StorageService

# =====================================================================
# helper functions for job progress logging
# =====================================================================

def log_job(job: Job, message: str, level: str = 'INFO'):
    JobLog.objects.create(job=job, message=message, level=level)

def update_progress(job: Job, progress: int):
    job.progress = progress
    job.save(update_fields=['progress', 'updated_at'])

# =====================================================================
# UNIFIED EXPORT WORKER PIPELINE
# =====================================================================

@shared_task(bind=True)
def run_export_pipeline_task(self, job_id: str):
    """
    Unified task processing PDF, DOCX, XLSX, or ZIP card record bundles.
    Configurations are loaded from the Job payload JSON field.
    """
    job = Job.objects.get(id=job_id)
    job.status = 'RUNNING'
    job.save()
    
    storage = StorageService()
    
    try:
        table_id = job.payload.get('table_id')
        export_type = job.payload.get('export_type')  # PDF, DOCX, XLSX, ZIP
        target_status = job.payload.get('status_filter')  # Optional status filter
        
        table = DynamicTable.objects.get(id=table_id)
        
        # Build query
        records = CardRecord.objects.filter(table=table)
        if target_status:
            records = records.filter(status=target_status)
        else:
            records = records.exclude(status='DELETED')
            
        total_records = records.count()
        if total_records == 0:
            raise ValueError("No records found matching export criteria.")
            
        log_job(job, f"Exporting {total_records} records in '{export_type}' format...")
        
        # Direct execution routing
        if export_type == 'XLSX':
            output_bytes, filename = generate_xlsx_export(table, records, job)
        elif export_type == 'ZIP':
            output_bytes, filename = generate_image_zip_export(table, records, job, storage)
        elif export_type == 'PDF':
            output_bytes, filename = generate_pdf_export(table, records, job, storage)
        elif export_type == 'DOCX':
            output_bytes, filename = generate_docx_export(table, records, job, storage)
        else:
            raise ValueError(f"Unsupported export type: {export_type}")
            
        # Upload output file to storage
        export_dir = f"exports/{table.id}/{export_type.lower()}/{uuid.uuid4()}_{filename}"
        result_url = storage.save_file(export_dir, io.BytesIO(output_bytes))
        
        # Update job outputs
        job.status = 'COMPLETED'
        job.result_url = result_url
        job.progress = 100
        job.save()
        log_job(job, f"Export completed successfully. Result URL: {result_url}")
        
    except Exception as e:
        job.status = 'FAILED'
        job.error_message = str(e)
        job.save()
        log_job(job, f"Fatal error during export generation: {str(e)}", 'ERROR')


# =====================================================================
# INDIVIDUAL FORMAT EXPORTERS
# =====================================================================

def generate_xlsx_export(table: DynamicTable, records: list, job: Job) -> (bytes, str):
    """Generates standard Excel file of dynamic fields."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Card Records"
    
    fields = table.fields.all()
    headers = [field.name for field in fields]
    ws.append(headers)
    
    total = len(records)
    for idx, record in enumerate(records):
        row = []
        for field in fields:
            if field.type == 'IMAGE':
                img_meta = record.images.get(field.key, {})
                row.append(img_meta.get('original_url', ''))
            else:
                row.append(record.data.get(field.key, ''))
        ws.append(row)
        
        # Update progress within 0-90% range
        progress = int((idx / total) * 90)
        update_progress(job, progress)
        
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue(), "data_export.xlsx"


def generate_image_zip_export(table: DynamicTable, records: list, job: Job, storage: StorageService) -> (bytes, str):
    """Zips all images within records."""
    output = io.BytesIO()
    total = len(records)
    
    with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for idx, record in enumerate(records):
            for img_key, meta in record.images.items():
                orig_path = meta.get('original_path')
                if orig_path:
                    try:
                        # Stream file from storage
                        file_content = storage.read_file(orig_path).read()
                        # Use lookup identifier or UUID as filename inside Zip
                        filename = f"{record.id}_{img_key}.jpg"
                        zip_file.writestr(filename, file_content)
                    except Exception as e:
                        log_job(job, f"Skipped file {orig_path} in export: {str(e)}", 'WARNING')
                        
            progress = int((idx / total) * 90)
            update_progress(job, progress)
            
    return output.getvalue(), "images_bundle.zip"


def generate_pdf_export(table: DynamicTable, records: list, job: Job, storage: StorageService) -> (bytes, str):
    """
    Generates a print-ready PDF using a custom layout.
    Uses reportlab or similar PDF canvas generators.
    """
    # reportlab PDF canvas generator
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    output = io.BytesIO()
    p = canvas.Canvas(output, pagesize=letter)
    width, height = letter
    
    total = len(records)
    for idx, record in enumerate(records):
        # Draw Card Frame
        p.rect(50, height - 350, 500, 300)
        p.drawString(100, height - 100, f"ID Card: {table.name}")
        p.drawString(100, height - 150, f"Record ID: {record.id}")
        
        y_offset = height - 180
        for key, val in record.data.items():
            p.drawString(100, y_offset, f"{key}: {val}")
            y_offset -= 20
            
        p.showPage()
        progress = int((idx / total) * 90)
        update_progress(job, progress)
        
    p.save()
    return output.getvalue(), "cards_print.pdf"


def generate_docx_export(table: DynamicTable, records: list, job: Job, storage: StorageService) -> (bytes, str):
    """
    Generates Word Document containing dynamic structures.
    Uses python-docx to structure layout tables.
    """
    import docx
    doc = docx.Document()
    doc.add_heading(f"ID Cards: {table.name}", 0)
    
    total = len(records)
    for idx, record in enumerate(records):
        doc.add_paragraph(f"Record ID: {record.id}")
        table_el = doc.add_table(rows=1, cols=2)
        hdr_cells = table_el.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Value'
        
        for key, val in record.data.items():
            row_cells = table_el.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(val)
            
        doc.add_page_break()
        progress = int((idx / total) * 90)
        update_progress(job, progress)
        
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue(), "cards_word.docx"
