"""
Word Export — Styles Mixin

Page setup, headers, footers, formatting, and XML helper methods
for Word document generation.
"""
from django.utils import timezone
from core.utils.template_rich_text import rich_text_to_plain_text


class WordStylesMixin:
    """Mixin providing page layout, header/footer, and cell styling methods."""

    @staticmethod
    def _is_hindi_abbasi_font(font_name_value) -> bool:
        """Return True when template font should use Abbasi."""
        val = str(font_name_value or '').strip().lower()
        return val in {'hindi', 'abbasi', 'abbasinatraj', 'abbasi_natraj'}

    @staticmethod
    def _force_run_font(run, font_name):
        """Force Word run font across ascii/hAnsi/eastAsia/complex scripts."""
        if not run or not font_name:
            return
        run.font.name = font_name
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass

    def _setup_page(self, doc, Cm, WD_ORIENT, parse_xml, nsdecls):
        """Configure page to landscape A4 with margins."""
        section = doc.sections[0]
        
        # Swap width and height for landscape
        new_width = section.page_height
        new_height = section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE
        
        # Set uniform margins to 0.5 cm on all sides.
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

        # Keep export width aligned with the actual section metrics so tables
        # respect side margins and avoid hard edge overflow.
        # python-docx subtraction yields raw EMU ints, not Length objects.
        # Convert EMU -> cm explicitly to avoid `'int' object has no attribute 'cm'`.
        page_w_emu = int(section.page_width)
        left_emu = int(section.left_margin)
        right_emu = int(section.right_margin)
        usable_width_cm = (page_w_emu - left_emu - right_emu) / 360000.0
        self.PAGE_WIDTH_CM = round(max(usable_width_cm, 1.0), 2)
    
    def _add_header(self, doc, institution_name, table_name, Cm, Pt, RGBColor,
                    WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls):
        """Add document header that repeats on EVERY page."""
        section = doc.sections[0]

        header = section.header
        header.is_linked_to_previous = False
        
        current_date = timezone.localtime(timezone.now()).strftime('%d-%m-%Y')
        
        # Create header table (2 cm narrower than page, centered)
        header_width = 25.5  # 27.5 - 2 = 25.5 cm
        header_table = header.add_table(rows=1, cols=3, width=Cm(header_width))
        header_table.autofit = False
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_cells = header_table.rows[0].cells
        
        # Set column widths (proportionally scaled to 25.5 cm)
        header_cells[0].width = Cm(8.5)
        header_cells[1].width = Cm(10)
        header_cells[2].width = Cm(7)
        
        # Left: Institution name
        left_para = header_cells[0].paragraphs[0]
        left_run = left_para.add_run(f'INSTITUTE NAME: {institution_name}')
        left_run.bold = True
        left_run.font.name = 'Arial'
        left_run.font.size = Pt(10)
        left_run.font.color.rgb = RGBColor(0, 0, 0)
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_para_spacing(left_para, parse_xml, nsdecls)
        
        # Center: Table name and date
        center_para = header_cells[1].paragraphs[0]
        center_run = center_para.add_run(f'{table_name} ({current_date})')
        center_run.bold = True
        center_run.font.name = 'Arial'
        center_run.font.size = Pt(11)
        center_run.font.color.rgb = RGBColor(0, 0, 0)
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_para_spacing(center_para, parse_xml, nsdecls)
        
        # Right: Branding
        right_para = header_cells[2].paragraphs[0]
        right_run = right_para.add_run('ADARSH ID CARDS')
        right_run.bold = True
        right_run.font.name = 'Arial'
        right_run.font.size = Pt(10)
        right_run.font.color.rgb = RGBColor(0, 0, 0)
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_para_spacing(right_para, parse_xml, nsdecls)
        
        # Remove borders and center vertically
        for cell in header_cells:
            self._remove_cell_borders(cell, parse_xml, nsdecls)
    
    def _add_footer(self, doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                    parse_xml, nsdecls, OxmlElement, qn, template_id=None):
        """Add document footer with selected-template text and page numbers."""
        from core.models import ExportTemplate
        
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        
        note_line = ''
        note_font_name = 'Arial'
        note_is_bold = False
        if template_id:
            try:
                tpl = ExportTemplate.objects.get(id=template_id)
                note_line = rich_text_to_plain_text(tpl.instructions or '').strip()
                note_font_name = 'AbbasiNatraj' if self._is_hindi_abbasi_font(tpl.font_name) else 'Arial'
                note_is_bold = bool(tpl.is_bold)
            except ExportTemplate.DoesNotExist:
                pass

        # Line 1: Selected template footer text (left side, can wrap up to multiple lines)
        if note_line:
            footer_para1 = footer.add_paragraph()
            lines = note_line.splitlines() or [note_line]
            for idx, line_text in enumerate(lines):
                if idx > 0:
                    footer_para1.add_run().add_break()
                footer_run = footer_para1.add_run(line_text)
                footer_run.bold = note_is_bold
                self._force_run_font(footer_run, note_font_name)
                footer_run.font.size = Pt(7)
                footer_run.font.color.rgb = RGBColor(0, 0, 0)
            footer_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._set_para_spacing(footer_para1, parse_xml, nsdecls, line=180)
        
        # Line 2: Page numbers on the right
        footer_para2 = footer.add_paragraph()
        self._set_para_spacing(footer_para2, parse_xml, nsdecls, line=180)
        
        # Add tab stop
        pPr = footer_para2._p.get_or_add_pPr()
        tabs = parse_xml(
            r'<w:tabs {}><w:tab w:val="right" w:pos="14400"/></w:tabs>'.format(nsdecls('w'))
        )
        pPr.append(tabs)
        
        # Tab
        footer_para2.add_run('\t')
        
        # Page X of Y
        self._add_page_numbers(footer_para2, Pt, RGBColor, OxmlElement, qn)
    
    def _add_page_numbers(self, para, Pt, RGBColor, OxmlElement, qn):
        """Add Page X of Y fields to paragraph."""
        # "Page " text
        page_run = para.add_run('Page ')
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(9)
        page_run.font.bold = True
        page_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # PAGE field
        self._add_field(para, 'PAGE', Pt, OxmlElement, qn)
        
        # " of " text
        of_run = para.add_run(' of ')
        of_run.font.name = 'Arial'
        of_run.font.size = Pt(9)
        of_run.font.bold = True
        of_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # NUMPAGES field
        self._add_field(para, 'NUMPAGES', Pt, OxmlElement, qn)
    
    def _add_field(self, para, field_name, Pt, OxmlElement, qn):
        """Add a Word field (PAGE, NUMPAGES, etc.) to paragraph."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_name
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run = para.add_run()
        run.font.size = Pt(9)
        run.font.bold = True
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
    
    def _set_compatibility_mode(self, doc):
        """Set document to Word 97-2003 compatibility mode.
        
        This tells Word to render the document in compatibility mode,
        which ensures maximum backward compatibility with older Word
        versions and enables VML image rendering.
        """
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            settings_elem = doc.settings.element
            
            # Find or create w:compat element
            compat = settings_elem.find(qn('w:compat'))
            if compat is None:
                compat = OxmlElement('w:compat')
                settings_elem.append(compat)
            
            # Remove existing compatibilityMode setting if any
            for cs in list(compat):
                if cs.get(qn('w:name')) == 'compatibilityMode':
                    compat.remove(cs)
            
            # Add Word 2003 compatibility mode (val=11)
            cs = OxmlElement('w:compatSetting')
            cs.set(qn('w:name'), 'compatibilityMode')
            cs.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
            cs.set(qn('w:val'), '11')
            compat.append(cs)
        except Exception:
            pass  # Not critical — VML images alone provide compatibility
    
    def _style_data_cell(self, cell, width, is_image, Cm, Pt, RGBColor,
                          WD_ALIGN_PARAGRAPH, parse_xml, nsdecls, font_pt=9):
        """Apply styling to a data cell with proper wrapping and padding."""
        if is_image:
            self._set_cell_margins(cell, parse_xml, nsdecls, 0, 0, 0, 0)
        else:
            # 1px-equivalent inner padding for denser exports (15 twips ~ 1px at 96dpi)
            self._set_cell_margins(cell, parse_xml, nsdecls, 15, 15, 15, 15)
        
        self._set_cell_vertical_align(cell, parse_xml, nsdecls)
        
        para = cell.paragraphs[0]
        if para.runs:
            run = para.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(font_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        self._set_para_spacing(para, parse_xml, nsdecls)
        cell.width = Cm(width)
        
        # Ensure cell allows text wrapping (remove noWrap flags)
        if not is_image:
            from lxml import etree
            from docx.oxml.ns import qn as _qn
            tcPr = cell._tc.get_or_add_tcPr()
            for nw in tcPr.findall(_qn('w:noWrap')):
                tcPr.remove(nw)
    
    def _set_table_borders(self, table, parse_xml, nsdecls):
        """Set table borders to 0.5pt."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            r'<w:tblPr {}/>'.format(nsdecls('w'))
        )
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
        
        # Remove existing borders
        for child in list(tblPr):
            if 'tblBorders' in child.tag:
                tblPr.remove(child)
        
        tblBorders = parse_xml(
            r'<w:tblBorders {}>'
            r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
            r'<w:left w:val="single" w:sz="4" w:color="000000"/>'
            r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
            r'<w:right w:val="single" w:sz="4" w:color="000000"/>'
            r'<w:insideH w:val="single" w:sz="4" w:color="000000"/>'
            r'<w:insideV w:val="single" w:sz="4" w:color="000000"/>'
            r'</w:tblBorders>'.format(nsdecls('w'))
        )
        tblPr.append(tblBorders)
    
    def _add_template_instructions(self, doc, template_id, Pt, RGBColor,
                                    WD_ALIGN_PARAGRAPH, parse_xml, nsdecls):
        """Add template instructions section after the data table."""
        from core.models import ExportTemplate
        try:
            tpl = ExportTemplate.objects.get(id=template_id)
        except ExportTemplate.DoesNotExist:
            return
        
        instructions = rich_text_to_plain_text(tpl.instructions or '').strip()
        if not instructions:
            return
        
        # Determine font from template settings
        use_hindi = self._is_hindi_abbasi_font(tpl.font_name)
        font_name = 'AbbasiNatraj' if use_hindi else 'Arial'
        use_bold = tpl.is_bold
        
        # Add blank line
        doc.add_paragraph('')
        
        # Instructions heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('INSTRUCTIONS:')
        heading_run.bold = True
        heading_run.underline = True
        self._force_run_font(heading_run, font_name)
        heading_run.font.size = Pt(9)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_para_spacing(heading_para, parse_xml, nsdecls, before=60, after=40)
        
        # Instructions body (preserve line breaks)
        for line in instructions.split('\n'):
            line = line.strip()
            if not line:
                continue
            body_para = doc.add_paragraph()
            body_run = body_para.add_run(line.upper())
            body_run.bold = use_bold
            self._force_run_font(body_run, font_name)
            body_run.font.size = Pt(8)
            body_run.font.color.rgb = RGBColor(0, 0, 0)
            body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._set_para_spacing(body_para, parse_xml, nsdecls, before=0, after=20, line=220)
    
    def _set_para_spacing(self, para, parse_xml, nsdecls, before=0, after=0, line=240):
        """Set paragraph spacing."""
        pPr = para._p.get_or_add_pPr()
        spacing = parse_xml(
            r'<w:spacing {} w:before="{}" w:after="{}" w:line="{}" w:lineRule="auto"/>'.format(
                nsdecls('w'), before, after, line
            )
        )
        pPr.append(spacing)
    
    def _set_cell_margins(self, cell, parse_xml, nsdecls, top=0, bottom=0, left=15, right=15):
        """Set cell margins in twips."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            r'<w:tcMar {}>'
            r'<w:top w:w="{}" w:type="dxa"/>'
            r'<w:bottom w:w="{}" w:type="dxa"/>'
            r'<w:left w:w="{}" w:type="dxa"/>'
            r'<w:right w:w="{}" w:type="dxa"/>'
            r'</w:tcMar>'.format(nsdecls('w'), top, bottom, left, right)
        )
        tcPr.append(tcMar)
    
    def _set_cell_vertical_align(self, cell, parse_xml, nsdecls, align='center'):
        """Set cell vertical alignment."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = parse_xml(r'<w:vAlign {} w:val="{}"/>'.format(nsdecls('w'), align))
        tcPr.append(vAlign)
    
    def _remove_cell_borders(self, cell, parse_xml, nsdecls):
        """Remove borders from a cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            r'<w:tcBorders {}>'
            r'<w:top w:val="nil"/>'
            r'<w:left w:val="nil"/>'
            r'<w:bottom w:val="nil"/>'
            r'<w:right w:val="nil"/>'
            r'</w:tcBorders>'.format(nsdecls('w'))
        )
        tcPr.append(tcBorders)
        vAlign = parse_xml(r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w')))
        tcPr.append(vAlign)
